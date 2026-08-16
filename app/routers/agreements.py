from __future__ import annotations

import json
import logging
import os
import shutil
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any, List, Optional

from fastapi import APIRouter, Depends, Form, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database import get_session
from app.models.models import Agreement, AgreementType, User
from app.routers.auth import get_user
from app.settings.settings import get_settings
from app.services.notification_service import NotificationService
from app.services.agreement_analyzer import MusicAgreementAnalyzer
from app.services.document_generator import DocumentGenerator

# Configure logging for re-extraction monitoring
logger = logging.getLogger(__name__)


# ============================================================================
# Persistent File Storage Utilities
# ============================================================================

# Storage directory for agreement files (relative to project root)
AGREEMENT_STORAGE_DIR = Path(__file__).parent.parent.parent / "storage" / "agreements"

# Maximum file size for re-extraction (32MB - matches Claude Vision limit)
MAX_FILE_SIZE_BYTES = 32 * 1024 * 1024


def get_storage_directory() -> Path:
    """Get and ensure the storage directory exists."""
    AGREEMENT_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    return AGREEMENT_STORAGE_DIR


def save_agreement_file(contents: bytes, filename: str) -> str:
    """
    Save agreement file to persistent storage.

    Args:
        contents: File contents as bytes
        filename: Unique filename (uuid + extension)

    Returns:
        str: Relative path from storage root (e.g., "agreements/uuid.pdf")
    """
    storage_dir = get_storage_directory()
    file_path = storage_dir / filename

    with open(file_path, "wb") as f:
        f.write(contents)

    # Return relative path for database storage
    return f"agreements/{filename}"


def get_agreement_file_path(storage_path: str) -> Path:
    """
    Get absolute path from stored relative path.

    Args:
        storage_path: Relative path from DB (e.g., "agreements/uuid.pdf")

    Returns:
        Path: Absolute path to file
    """
    # storage_path is "agreements/filename", storage_dir is already /storage/agreements
    # So we need to go up one level
    storage_root = AGREEMENT_STORAGE_DIR.parent
    return storage_root / storage_path


def delete_agreement_file(storage_path: str) -> bool:
    """
    Delete agreement file from persistent storage.

    Args:
        storage_path: Relative path from DB

    Returns:
        bool: True if deleted, False if not found
    """
    if not storage_path:
        return False

    try:
        file_path = get_agreement_file_path(storage_path)
        if file_path.exists():
            file_path.unlink()
            logger.info(f"Deleted agreement file: {storage_path}")
            return True
    except Exception as e:
        logger.warning(f"Failed to delete agreement file {storage_path}: {e}")
    return False


def validate_file_for_reextraction(file_path: Path, file_size: int) -> dict:
    """
    Validate file is suitable for re-extraction.

    Args:
        file_path: Path to the file
        file_size: Expected file size

    Returns:
        dict: {valid: bool, error: str|None}
    """
    if not file_path.exists():
        return {"valid": False, "error": "Original file not found in storage"}

    actual_size = file_path.stat().st_size
    if actual_size != file_size:
        logger.warning(f"File size mismatch: expected {file_size}, got {actual_size}")

    if actual_size > MAX_FILE_SIZE_BYTES:
        return {"valid": False, "error": f"File too large for re-extraction (max {MAX_FILE_SIZE_BYTES // (1024*1024)}MB)"}

    return {"valid": True, "error": None}

agreements_router = APIRouter(
    prefix="/agreements",
    tags=["Agreements"],
)


class AgreementResponse(BaseModel):
    id: int
    filename: str
    original_filename: str
    file_size: int
    agreement_type: str
    uploaded_at: str
    # Analysis summary fields (optional - only present if parsed_content exists)
    overall_rating: Optional[str] = None  # "FAVORABLE", "NEUTRAL", "UNFAVORABLE"
    red_flag_count: Optional[int] = None  # Total red flags detected
    critical_flags_count: Optional[int] = None  # CRITICAL severity flags
    high_flags_count: Optional[int] = None  # HIGH severity flags
    medium_flags_count: Optional[int] = None  # MEDIUM severity flags
    # Extraction quality fields
    extraction_quality_score: Optional[int] = None  # 0-100 quality score
    extraction_method: Optional[str] = None  # "vision_api", "pdfplumber", etc.

    class Config:
        from_attributes = True


class AgreementListResponse(BaseModel):
    agreements: List[AgreementResponse]
    total: int


class AgreementDetailResponse(BaseModel):
    """
    Detailed agreement response including parsed analysis.

    The `parsed_content` field contains the comprehensive analysis from
    MusicAgreementAnalyzer for producer agreements, including:
    - meta: Document hash and analysis date for verification
    - agreement: Type, parties, effective date, context
    - overall_assessment: Rating with color-coded counts and summary
    - terms: Financial, rights, credit, legal, administrative, publishing sections
    - red_flags: Array of detected issues (RF01-RF18) with severity and recommendations
    - negotiation_priorities: Ranked list of terms to negotiate
    - financial_projection: Scenario-based financial forecasting

    The `analysis_version` field tracks which analyzer version produced the results,
    enabling future schema migrations and backward compatibility.
    """
    id: int
    filename: str
    original_filename: str
    file_size: int
    agreement_type: str
    uploaded_at: str
    parsed_content: Optional[EnhancedAgreementAnalysis] = None
    analysis_version: Optional[str] = None  # Analyzer version (e.g., "1.0")

    class Config:
        from_attributes = True


class AgreementUpdateRequest(BaseModel):
    parsed_content: Optional[EnhancedAgreementAnalysis] = None
    agreement_type: Optional[str] = None


# ============================================================================
# Enhanced Analysis Schema Models
# These provide type safety and documentation for the MusicAgreementAnalyzer output
# ============================================================================

from enum import Enum as PyEnum


class TermColor(str, PyEnum):
    """Color coding for term assessments relative to industry standards."""
    RED = "RED"       # Unfavorable - worse than industry standard
    YELLOW = "YELLOW" # Neutral - matches industry standard
    GREEN = "GREEN"   # Favorable - better than industry standard
    GRAY = "GRAY"     # Not found/specified in document


class RedFlagSeverity(str, PyEnum):
    """Severity levels for detected red flags (RF01-RF18)."""
    CRITICAL = "CRITICAL"  # Major financial/legal risk
    HIGH = "HIGH"          # Significant concern
    MEDIUM = "MEDIUM"      # Notable issue


class TermDetail(BaseModel):
    """
    Individual term with color-coded assessment.

    Example:
        {
            "value": "2% of Net PPD",
            "clause": "Section 4.1",
            "color": "YELLOW",
            "assessment": "Standard producer royalty rate",
            "industry_standard": "3-5% direct"
        }
    """
    value: Optional[str] = None
    clause: Optional[str] = None
    color: Optional[TermColor] = None
    assessment: Optional[str] = None
    industry_standard: Optional[str] = None

    class Config:
        use_enum_values = True


class FinancialTerms(BaseModel):
    """Financial terms analysis (advance, royalties, sync share, etc.)."""
    nominal_fee: Optional[TermDetail] = None
    advance: Optional[TermDetail] = None
    recoupable_advance: Optional[TermDetail] = None
    royalty_rate: Optional[TermDetail] = None
    royalty_base: Optional[TermDetail] = None
    payment_threshold: Optional[TermDetail] = None
    sync_share: Optional[TermDetail] = None
    escalation: Optional[TermDetail] = None
    recoupment: Optional[TermDetail] = None
    sync_license_share: Optional[TermDetail] = None


class RightsTerms(BaseModel):
    """Rights granted analysis (duration, territory, media scope, etc.)."""
    grant_type: Optional[TermDetail] = None
    duration: Optional[TermDetail] = None
    territory: Optional[TermDetail] = None
    media_scope: Optional[TermDetail] = None
    remix_rights: Optional[TermDetail] = None
    name_likeness: Optional[TermDetail] = None
    touring_visuals: Optional[TermDetail] = None
    derivative_works: Optional[TermDetail] = None


class CreditTerms(BaseModel):
    """Credit and attribution terms analysis."""
    credit_format: Optional[TermDetail] = None
    credit_placement: Optional[TermDetail] = None
    credit_remedy: Optional[TermDetail] = None
    likeness_approval: Optional[TermDetail] = None


class LegalTerms(BaseModel):
    """Legal protections analysis (warranties, audit rights, indemnification, etc.)."""
    warranties: Optional[TermDetail] = None
    indemnification: Optional[TermDetail] = None
    moral_rights: Optional[TermDetail] = None
    third_party_payments: Optional[TermDetail] = None
    audit_rights: Optional[TermDetail] = None
    objection_period: Optional[TermDetail] = None
    litigation_deadline: Optional[TermDetail] = None
    dispute_resolution: Optional[TermDetail] = None


class AdministrativeTerms(BaseModel):
    """Administrative terms analysis (accounting, notices, amendments, etc.)."""
    accounting_frequency: Optional[TermDetail] = None
    payment_timing: Optional[TermDetail] = None
    notices: Optional[TermDetail] = None
    amendment: Optional[TermDetail] = None
    assignment_rights: Optional[TermDetail] = None
    governing_law: Optional[TermDetail] = None


class PublishingTerms(BaseModel):
    """Publishing-specific terms analysis (if applicable)."""
    composition_ownership: Optional[TermDetail] = None
    controlled_composition: Optional[TermDetail] = None
    sync_license: Optional[TermDetail] = None
    mechanical_royalties: Optional[TermDetail] = None
    performance_royalties: Optional[TermDetail] = None
    print_rights: Optional[TermDetail] = None


class TermsAnalysis(BaseModel):
    """Container for all term categories."""
    financial: Optional[FinancialTerms] = None
    rights: Optional[RightsTerms] = None
    credit: Optional[CreditTerms] = None
    legal: Optional[LegalTerms] = None
    administrative: Optional[AdministrativeTerms] = None
    publishing: Optional[PublishingTerms] = None


class RedFlag(BaseModel):
    """
    Critical issue detected in agreement (RF01-RF18).

    Example:
        {
            "id": "RF01",
            "name": "Fraction/Formula Royalty",
            "severity": "CRITICAL",
            "clause": "Section 4.2",
            "quote": "royalty shall be calculated as Producer Rate divided by Artist Rate",
            "impact": "Effective rate becomes ~10% of headline rate",
            "recommendation": "Negotiate for direct percentage royalty instead of fraction formula"
        }
    """
    id: Optional[str] = None
    name: Optional[str] = None
    severity: Optional[RedFlagSeverity] = None
    clause: Optional[str] = None
    quote: Optional[str] = None
    trigger: Optional[str] = None
    impact: Optional[str] = None
    recommendation: Optional[str] = None

    class Config:
        use_enum_values = True


class NegotiationPriority(BaseModel):
    """
    Ranked item for negotiation based on impact and achievability.

    Example:
        {
            "priority": 1,
            "term": "royalty_rate",
            "issue": "Below industry standard",
            "current": "2% via fraction",
            "target": "3-5% direct",
            "impact": "HIGH",
            "achievability": "MEDIUM"
        }
    """
    priority: Optional[int] = None
    term: Optional[str] = None
    issue: Optional[str] = None
    current: Optional[str] = None
    target: Optional[str] = None
    impact: Optional[str] = None
    achievability: Optional[str] = None


class FinancialProjection(BaseModel):
    """
    Financial forecasting based on agreement terms.

    Example:
        {
            "scenario": "moderate_success",
            "estimated_advance": "$5,000",
            "estimated_recording_royalties": "$2,500/year",
            "estimated_sync_income": "$1,000/year",
            "key_insight": "Recoupment expected within 2 years"
        }
    """
    scenario: Optional[str] = None
    estimated_advance: Optional[str] = None
    estimated_recording_royalties: Optional[str] = None
    estimated_sync_income: Optional[str] = None
    estimated_total_annual: Optional[str] = None
    key_insight: Optional[str] = None


class OverallAssessment(BaseModel):
    """
    Summary metrics and rating for the agreement.

    Example:
        {
            "rating": "UNFAVORABLE",
            "red_count": 3,
            "yellow_count": 8,
            "green_count": 4,
            "gray_count": 2,
            "critical_flags": 1,
            "high_flags": 2,
            "medium_flags": 0,
            "summary": "Agreement contains significant red flags..."
        }
    """
    rating: Optional[str] = None
    red_count: Optional[int] = None
    yellow_count: Optional[int] = None
    green_count: Optional[int] = None
    gray_count: Optional[int] = None
    critical_flags: Optional[int] = None
    high_flags: Optional[int] = None
    medium_flags: Optional[int] = None
    summary: Optional[str] = None


class AgreementParty(BaseModel):
    """Party information extracted from agreement."""
    name: Optional[str] = None
    role: Optional[str] = None
    address: Optional[str] = None


class AgreementParties(BaseModel):
    """All parties involved in the agreement."""
    party_a: Optional[AgreementParty] = None
    party_b: Optional[AgreementParty] = None
    artist: Optional[str] = None
    distributor: Optional[str] = None
    track_or_project: Optional[str] = None


class AgreementInfo(BaseModel):
    """
    Agreement metadata extracted from document.

    Example:
        {
            "type": "SAMPLE_CLEARANCE",
            "parties": {...},
            "effective_date": "2024-01-15",
            "context": "Sample clearance for use in new recording"
        }
    """
    type: Optional[str] = None
    parties: Optional[AgreementParties] = None
    effective_date: Optional[str] = None
    context: Optional[Any] = None  # Can be string or dict depending on analyzer version


class AgreementMeta(BaseModel):
    """Metadata about the analysis itself."""
    document_hash: Optional[str] = None
    analysis_date: Optional[str] = None


class EnhancedAgreementAnalysis(BaseModel):
    """
    Complete analysis structure from MusicAgreementAnalyzer.

    This is the comprehensive output format for producer agreements analyzed
    with the 18-red-flag detection system.
    """
    meta: Optional[AgreementMeta] = None
    agreement: Optional[AgreementInfo] = None
    overall_assessment: Optional[OverallAssessment] = None
    terms: Optional[TermsAnalysis] = None
    red_flags: Optional[List[RedFlag]] = None
    negotiation_priorities: Optional[List[NegotiationPriority]] = None
    financial_projection: Optional[FinancialProjection] = None

    # Backward compatibility fields
    agreement_summary: Optional[dict] = None
    overall_score: Optional[dict] = None
    royalty_analysis: Optional[dict] = None
    field_ratings: Optional[dict] = None
    general_assessment: Optional[str] = None
    text_preview: Optional[str] = None
    full_text: Optional[str] = None


class ExtractionResult:
    """Container for text extraction results with metadata"""
    def __init__(self, text: str, method: str, quality_score: int = 0,
                 warnings: List[str] = None, page_count: int = 0):
        self.text = text
        self.method = method  # "vision_api", "pdfplumber", "pypdf2", "docx"
        self.quality_score = quality_score  # 0-100
        self.warnings = warnings or []
        self.page_count = page_count
        self.character_count = len(text)


def validate_extraction(text: str, source_type: str = "pdf") -> dict:
    """
    Validate extraction quality and return a quality report.

    Returns:
        dict with quality_score (0-100), warnings, and is_valid
    """
    warnings = []
    text_lower = text.lower()

    # Calculate quality score (0-100)
    score = 100
    char_count = len(text.strip())

    # Penalize short extractions more heavily
    if char_count < 500:
        score -= 60
        warnings.append("Very short extraction (<500 chars) - document may not have extracted properly")
    elif char_count < 1000:
        score -= 40
        warnings.append("Short extraction (<1000 chars) - may be missing content")
    elif char_count < 2000:
        score -= 20
        warnings.append("Moderate extraction length - verify completeness")

    # Reward presence of agreement keywords (more comprehensive)
    agreement_keywords = [
        "agreement", "contract", "party", "parties", "whereas", "witnesseth",
        "royalty", "royalties", "advance", "recoup", "territory", "term",
        "grant", "license", "rights", "obligations", "payment", "compensation"
    ]
    keyword_count = sum(1 for kw in agreement_keywords if kw in text_lower)

    if keyword_count < 3:
        score -= 30
        warnings.append(f"Few agreement keywords found ({keyword_count}/18) - may not be a valid agreement")
    elif keyword_count < 6:
        score -= 15
        warnings.append(f"Limited agreement keywords ({keyword_count}/18) - extraction may be incomplete")
    elif keyword_count >= 12:
        score += 10  # Bonus for comprehensive extraction

    # Check for encoding issues (more strict)
    encoding_issues = text.count("�") + text.count("\\x")
    if encoding_issues > 20:
        score -= 30
        warnings.append(f"Significant encoding issues detected ({encoding_issues} instances)")
    elif encoding_issues > 5:
        score -= 15
        warnings.append(f"Some encoding issues detected ({encoding_issues} instances)")

    # Check for excessive whitespace (sign of poor extraction)
    whitespace_ratio = (text.count("\n\n\n") + text.count("   ")) / max(len(text), 1)
    if whitespace_ratio > 0.05:
        score -= 20
        warnings.append("Excessive whitespace detected - formatting may be corrupted")

    # Check for table markers (good sign)
    if "TABLE" in text or "|" in text:
        score += 5

    # Check for section markers (good sign)
    if any(marker in text for marker in ["Section", "Article", "Clause", "Exhibit", "Schedule"]):
        score += 5

    # Ensure score stays in 0-100 range
    score = max(0, min(100, score))

    return {
        "quality_score": score,
        "warnings": warnings,
        "is_valid": score >= 30,
        "character_count": char_count,
        "keyword_count": keyword_count
    }


def extract_text_from_pdf_vision(file_path: str) -> ExtractionResult:
    """
    Extract text from PDF using Claude's PDF Vision API.
    Best for scanned PDFs, complex layouts, and tables.

    Returns:
        ExtractionResult with extracted text and metadata
    """
    import anthropic
    import base64
    import os

    settings = get_settings()

    try:
        # Check file size (Claude PDF Vision limit is 32MB)
        file_size = os.path.getsize(file_path)
        if file_size > 32 * 1024 * 1024:
            return ExtractionResult(
                text="",
                method="vision_api",
                quality_score=0,
                warnings=["File too large for PDF Vision API (>32MB)"]
            )

        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

        with open(file_path, "rb") as pdf_file:
            pdf_data = base64.standard_b64encode(pdf_file.read()).decode("utf-8")

        extraction_prompt = """Extract ALL text from this document, preserving the structure as much as possible.

IMPORTANT INSTRUCTIONS:
1. Extract ALL text including headers, footers, tables, and footnotes
2. For tables, format as:
   --- TABLE ---
   [Row 1: Col1 | Col2 | Col3]
   [Row 2: Col1 | Col2 | Col3]
   --- END TABLE ---
3. Preserve paragraph breaks with blank lines
4. Include any signatures, dates, and page numbers
5. Mark section headers clearly
6. Do NOT summarize or interpret - extract verbatim text only

Return ONLY the extracted text, no commentary."""

        message_content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_data,
                },
            },
            {"type": "text", "text": extraction_prompt}
        ]

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,
            messages=[{"role": "user", "content": message_content}]
        )

        extracted_text = response.content[0].text if response.content else ""

        # Validate extraction
        validation = validate_extraction(extracted_text, "pdf")

        print(f"[PDF Vision] Extracted {len(extracted_text)} characters, quality: {validation['quality_score']}")

        return ExtractionResult(
            text=extracted_text,
            method="vision_api",
            quality_score=validation["quality_score"],
            warnings=validation["warnings"]
        )

    except anthropic.APIError as e:
        print(f"[PDF Vision] API error: {e}")
        return ExtractionResult(
            text="",
            method="vision_api",
            quality_score=0,
            warnings=[f"PDF Vision API error: {str(e)}"]
        )
    except Exception as e:
        print(f"[PDF Vision] Error: {e}")
        return ExtractionResult(
            text="",
            method="vision_api",
            quality_score=0,
            warnings=[f"PDF Vision extraction failed: {str(e)}"]
        )


def extract_text_from_pdf_pdfplumber(file_path: str) -> ExtractionResult:
    """
    Extract text from PDF using pdfplumber.
    Better than PyPDF2 for tables and complex layouts.

    Returns:
        ExtractionResult with extracted text and metadata
    """
    try:
        import pdfplumber

        all_text = []
        page_count = 0
        table_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                page_text = []

                # Extract tables first
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        table_count += 1
                        page_text.append(f"\n--- TABLE {table_count} ---")
                        for row in table:
                            # Filter out None values and join cells
                            row_text = " | ".join(str(cell) if cell else "" for cell in row)
                            page_text.append(f"[{row_text}]")
                        page_text.append("--- END TABLE ---\n")

                # Extract regular text
                text = page.extract_text() or ""
                if text.strip():
                    page_text.append(text)

                if page_text:
                    all_text.append(f"--- PAGE {i+1} ---\n" + "\n".join(page_text))

        combined_text = "\n\n".join(all_text)

        # Validate extraction
        validation = validate_extraction(combined_text, "pdf")

        warnings = validation["warnings"]
        if table_count > 0:
            print(f"[pdfplumber] Extracted {table_count} tables from {page_count} pages")

        print(f"[pdfplumber] Extracted {len(combined_text)} characters, quality: {validation['quality_score']}")

        return ExtractionResult(
            text=combined_text,
            method="pdfplumber",
            quality_score=validation["quality_score"],
            warnings=warnings,
            page_count=page_count
        )

    except Exception as e:
        print(f"[pdfplumber] Error: {e}")
        return ExtractionResult(
            text="",
            method="pdfplumber",
            quality_score=0,
            warnings=[f"pdfplumber extraction failed: {str(e)}"]
        )


def extract_text_from_pdf_pypdf2(file_path: str) -> ExtractionResult:
    """
    Extract text from PDF using PyPDF2 (fallback method).
    Basic text extraction, works for simple PDFs.

    Returns:
        ExtractionResult with extracted text and metadata
    """
    try:
        import PyPDF2

        text_parts = []
        page_count = 0

        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"--- PAGE {i+1} ---\n{page_text}")

        combined_text = "\n\n".join(text_parts)

        # Validate extraction
        validation = validate_extraction(combined_text, "pdf")

        print(f"[PyPDF2] Extracted {len(combined_text)} characters, quality: {validation['quality_score']}")

        return ExtractionResult(
            text=combined_text,
            method="pypdf2",
            quality_score=validation["quality_score"],
            warnings=validation["warnings"],
            page_count=page_count
        )

    except Exception as e:
        print(f"[PyPDF2] Error: {e}")
        return ExtractionResult(
            text="",
            method="pypdf2",
            quality_score=0,
            warnings=[f"PyPDF2 extraction failed: {str(e)}"]
        )


def extract_text_from_pdf(file_path: str, preferred_method: str = "auto") -> ExtractionResult:
    """
    Extract text from PDF using multi-layer strategy.

    Strategy:
    - auto: Try pdfplumber first, then Vision API if quality < 50, then PyPDF2
    - vision: Force Claude PDF Vision API
    - standard: Use pdfplumber/pypdf2 only (no API calls)

    Args:
        file_path: Path to PDF file
        preferred_method: "auto", "vision", or "standard"

    Returns:
        ExtractionResult with best extraction result
    """
    import os

    print(f"[PDF Extraction] Starting extraction for {os.path.basename(file_path)}, method: {preferred_method}")

    # Force vision extraction
    if preferred_method == "vision":
        result = extract_text_from_pdf_vision(file_path)
        if result.quality_score >= 30:
            return result
        # Fall back to pdfplumber if vision fails
        print("[PDF Extraction] Vision API failed, falling back to pdfplumber")
        result = extract_text_from_pdf_pdfplumber(file_path)
        if result.quality_score >= 30:
            return result
        return extract_text_from_pdf_pypdf2(file_path)

    # Standard extraction (no API calls)
    if preferred_method == "standard":
        result = extract_text_from_pdf_pdfplumber(file_path)
        if result.quality_score >= 30:
            return result
        return extract_text_from_pdf_pypdf2(file_path)

    # Auto strategy: pdfplumber -> Vision API (if needed) -> PyPDF2
    result = extract_text_from_pdf_pdfplumber(file_path)

    if result.quality_score >= 50:
        print(f"[PDF Extraction] pdfplumber succeeded with quality {result.quality_score}")
        return result

    # Try Vision API for better extraction
    print(f"[PDF Extraction] pdfplumber quality {result.quality_score} < 50, trying Vision API")
    vision_result = extract_text_from_pdf_vision(file_path)

    if vision_result.quality_score > result.quality_score:
        print(f"[PDF Extraction] Vision API improved quality: {vision_result.quality_score}")
        return vision_result

    # If pdfplumber was decent, use it
    if result.quality_score >= 30:
        print(f"[PDF Extraction] Using pdfplumber result (quality {result.quality_score})")
        return result

    # Last resort: try PyPDF2
    print("[PDF Extraction] Trying PyPDF2 as fallback")
    pypdf2_result = extract_text_from_pdf_pypdf2(file_path)

    # Return best result
    best = max([result, vision_result, pypdf2_result], key=lambda r: r.quality_score)
    print(f"[PDF Extraction] Best result: {best.method} with quality {best.quality_score}")
    return best


def extract_text_from_docx(file_path: str) -> ExtractionResult:
    """
    Extract text from DOCX file with comprehensive extraction.
    Extracts paragraphs, tables, headers, and footers.

    Returns:
        ExtractionResult with extracted text and metadata
    """
    try:
        import docx

        doc = docx.Document(file_path)
        all_text = ["--- DOCUMENT START ---"]
        table_count = 0

        # Extract main body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_count += 1
            all_text.append(f"\n--- TABLE {table_count} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                all_text.append(f"[{row_text}]")
            all_text.append("--- END TABLE ---\n")

        # Extract headers and footers
        header_text = []
        footer_text = []
        for section in doc.sections:
            # Headers
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text)
            # Footers
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text)

        if header_text:
            all_text.insert(1, "\n--- HEADER ---\n" + "\n".join(set(header_text)) + "\n--- END HEADER ---\n")

        if footer_text:
            all_text.append("\n--- FOOTER ---\n" + "\n".join(set(footer_text)) + "\n--- END FOOTER ---")

        all_text.append("\n--- DOCUMENT END ---")

        combined_text = "\n".join(all_text)

        # Validate extraction
        validation = validate_extraction(combined_text, "docx")

        print(f"[DOCX] Extracted {len(combined_text)} characters, {table_count} tables, quality: {validation['quality_score']}")

        return ExtractionResult(
            text=combined_text,
            method="docx",
            quality_score=validation["quality_score"],
            warnings=validation["warnings"]
        )

    except Exception as e:
        print(f"[DOCX] Error: {e}")
        return ExtractionResult(
            text="",
            method="docx",
            quality_score=0,
            warnings=[f"DOCX extraction failed: {str(e)}"]
        )


def detect_agreement_type(text: str, filename: str, file_path: str = None) -> AgreementType:
    """
    Analyze document content with AI to detect agreement type.
    Uses Claude to analyze the actual content of the agreement.
    Falls back to keyword matching if AI fails.
    """
    import anthropic
    import base64

    settings = get_settings()

    # Try AI-based detection first (analyzes actual content)
    try:
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

        # Prepare content for Claude - optimized for Haiku
        classification_prompt = """TASK: Classify this music industry agreement into exactly ONE category.

CATEGORIES:
- "producer agreement": Producer/artist deals for beats, master recordings, producer points, work-for-hire, sample clearances
- "publishing": Songwriter/publisher deals for compositions, mechanical royalties, sync licenses, administration
- "management": Artist/manager deals for career representation, management commission

CLASSIFICATION RULES:
1. If document mentions "producer", "beat", "master recording", "sample", "points on masters" → producer agreement
2. If document mentions "songwriter", "composition", "publishing", "mechanical royalties", "sync" → publishing
3. If document mentions "manager", "management services", "career", "commission on earnings" → management

OUTPUT: Reply with ONLY one word (no explanation, no quotes):
producer agreement
OR
publishing
OR
management"""

        # Check if we should use PDF vision (if text is short and we have a PDF)
        use_pdf_vision = False
        if file_path and file_path.lower().endswith('.pdf') and len(text.strip()) < 500:
            use_pdf_vision = True

        if use_pdf_vision and file_path:
            # Use PDF vision for better analysis
            with open(file_path, "rb") as pdf_file:
                pdf_data = base64.standard_b64encode(pdf_file.read()).decode("utf-8")

            message_content = [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_data,
                    },
                },
                {"type": "text", "text": classification_prompt}
            ]

            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=50,
                messages=[{"role": "user", "content": message_content}]
            )
        else:
            # Use text-based analysis
            # Take first 8000 chars for classification (enough to understand document type)
            text_sample = text[:8000] if len(text) > 8000 else text

            response = client.messages.create(
                model="claude-3-5-haiku-20241022",
                max_tokens=20,
                messages=[{
                    "role": "user",
                    "content": f"{classification_prompt}\n\n--- DOCUMENT ---\n{text_sample}"
                }]
            )

        result = response.content[0].text.strip().lower()
        print(f"[DEBUG] AI classification result: '{result}'")

        if "producer" in result:
            print(f"[DEBUG] AI detected: PRODUCER AGREEMENT")
            return AgreementType.PRODUCER_AGREEMENT
        elif "management" in result:
            print(f"[DEBUG] AI detected: MANAGEMENT")
            return AgreementType.MANAGEMENT
        elif "publishing" in result:
            print(f"[DEBUG] AI detected: PUBLISHING")
            return AgreementType.PUBLISHING
        else:
            print(f"[DEBUG] AI returned unexpected result: '{result}', falling back to keyword matching")

    except Exception as e:
        print(f"[DEBUG] AI classification failed: {e}, falling back to keyword matching")

    # Fallback: keyword-based detection
    text_lower = text.lower()

    producer_keywords = [
        "producer agreement", "producer shall", "producer will", "producer hereby",
        "producer's services", "producer royalty", "producer fee", "producer points",
        "work for hire", "work-for-hire", "beat lease", "beat license", "beat sale",
        "master recording", "master side", "sound recording", "production services",
        "backend points", "points on the master", "all-in royalty", "producer advance"
    ]

    publishing_keywords = [
        "publishing agreement", "music publishing", "songwriter agreement",
        "mechanical royalt", "performance royalt", "synchronization", "sync license",
        "co-publishing", "sub-publishing", "copyright assignment", "writer share",
        "publisher share", "musical composition", "songwriter", "publishing rights",
        "administration agreement", "publishing administration"
    ]

    management_keywords = [
        "management agreement", "artist management", "talent management",
        "personal manager", "business manager", "manager commission",
        "management services", "artist representation", "career guidance"
    ]

    producer_score = sum(1 for kw in producer_keywords if kw in text_lower)
    publishing_score = sum(1 for kw in publishing_keywords if kw in text_lower)
    management_score = sum(1 for kw in management_keywords if kw in text_lower)

    print(f"[DEBUG] Keyword fallback - Producer: {producer_score}, Publishing: {publishing_score}, Management: {management_score}")

    scores = {
        AgreementType.PRODUCER_AGREEMENT: producer_score,
        AgreementType.PUBLISHING: publishing_score,
        AgreementType.MANAGEMENT: management_score,
    }

    max_score = max(scores.values())
    if max_score == 0:
        return AgreementType.PUBLISHING  # Default

    return max(scores, key=scores.get)


def get_producer_agreement_prompt(text: str) -> str:
    """Generate AI prompt for producer/sample agreement parsing - comprehensive analyzer"""
    # Generate a document hash for verification
    doc_hash = text[:50].replace('\n', ' ').replace('"', "'") if text else "EMPTY"

    return f"""You are a music contract analyst. You analyze ONLY the agreement text provided in the current message.

CRITICAL RULES:
1. Extract terms ONLY from the document provided — never invent or assume terms
2. If a term is not explicitly stated in the document, mark it as "Not specified"
3. Do NOT use knowledge from other agreements or examples
4. Each analysis is independent — ignore any prior conversation context
5. Quote the relevant clause number when extracting each term

Document hash for verification: "{doc_hash}"

## STEP 1: IDENTIFY AGREEMENT TYPE

First, determine the agreement type by looking for these indicators:

| Type | Key Indicators |
|------|----------------|
| sample_clearance | "sample", "sampled recording", "clearance", "new recording embodying" |
| work_for_hire | "work for hire", "work made for hire", "all right title and interest" |
| points_deal | "producer royalty", "points", "percentage of net receipts" |
| beat_lease | "lease", "non-exclusive license", "limited license" |
| beat_sale | "exclusive", "full buyout", "transfer of ownership" |
| co_production | "co-production", "joint production", "shared ownership" |

## STEP 2: EXTRACT PARTIES

Extract exactly as written in document:
- licensor: The party granting rights (producer/sample owner)
- licensee: The party receiving rights (label/artist)
- artist: Name of performing artist (if different from parties)
- track: Name of recording/composition/project

## STEP 3: EXTRACT AND ASSESS TERMS

### COLOR CODING RULES (IMPORTANT)
- **GREEN** = FAVORABLE - Term is BETTER than industry standard (benefits the licensor/producer)
- **YELLOW** = INDUSTRY STANDARD - Term matches typical industry practice (neither good nor bad)
- **RED** = UNFAVORABLE - Term is WORSE than industry standard (disadvantages the licensor/producer)

For EACH term below:
1. Search the document for the relevant clause
2. If found: extract exact value and cite clause number
3. If NOT found: set value to "Not specified" and color to "YELLOW"
4. Compare to industry standard and assign appropriate color
5. NEVER guess or infer values not explicitly stated

### FINANCIAL TERMS - Industry Standards
Note: If neither RED nor GREEN condition applies, use **YELLOW** (industry standard)
Note: Only include nominal_fee if explicitly mentioned in document (common in UK/EU, not US)

| Term | Field Name | Industry Standard (YELLOW) | RED if (worse) | GREEN if (better) |
|------|------------|----------------------------|----------------|-------------------|
| Nominal Fee | nominal_fee | £1-100 token (ONLY if mentioned) | N/A | N/A |
| Advance | advance | £3,000-10,000 | < £3,000 | > £10,000 |
| Recoupable Advance | recoupable_advance | Fully recoupable (standard) | Cross-collateralized with other projects | Non-recoupable or capped recoupment |
| Royalty Rate | royalty_rate | 3-5% direct | < 3% OR via Applicable Fraction | > 5% direct |
| Royalty Base | royalty_base | Net PPD with defined deductions | Undefined deductions OR references unseen agreement | Capped deductions specified |
| Payment Threshold | payment_threshold | £50-100 | > £200 | < £50 |
| Third-Party Sync | sync_share | 15-25% of label's receipts | < 15% OR via Applicable Fraction | > 25% direct |
| Escalation | escalation | Standard sales increases | Explicitly excluded | Included with specific milestones |

### RIGHTS GRANTED - Industry Standards
Note: If neither RED nor GREEN condition applies, use **YELLOW** (industry standard)

| Term | Field Name | Industry Standard (YELLOW) | RED if (worse) | GREEN if (better) |
|------|------------|----------------------------|----------------|-------------------|
| Duration | duration | Perpetuity (standard) | N/A | Reversion clause included |
| Territory | territory | Worldwide | N/A | Limited territory |
| Media | media_scope | All media | Unknown future formats without limit | Specific formats listed |
| Remixes | remix_rights | 3-5 included | Unlimited remixes | Capped or additional fee required |
| Name/Likeness | name_likeness | Credit + marketing use | Broad exploitation without limits | Limited use specified |
| Touring Visuals | touring_visuals | Separate negotiation | Royalty-free perpetual | Fee included or excluded |

### CREDIT & ATTRIBUTION - Industry Standards
Note: If neither RED nor GREEN condition applies, use **YELLOW** (industry standard)

| Term | Field Name | Industry Standard (YELLOW) | RED if (worse) | GREEN if (better) |
|------|------------|----------------------------|----------------|-------------------|
| Credit Format | credit_format | Sample credit in metadata | No credit specified | Specific format guaranteed |
| Credit Placement | credit_placement | Packaging + metadata | "Where possible" qualifier | All formats guaranteed |
| Failure Remedy | credit_remedy | Prospective cure | No remedy | Damages available |

### LEGAL PROTECTIONS - Industry Standards
Note: If neither RED nor GREEN condition applies, use **YELLOW** (industry standard)

| Term | Field Name | Industry Standard (YELLOW) | RED if (worse) | GREEN if (better) |
|------|------------|----------------------------|----------------|-------------------|
| Warranties | warranties | Standard warranties | Unusually broad | Limited/reasonable scope |
| Indemnification | indemnification | Mutual indemnification | One-sided + unlimited | Mutual + capped |
| Moral Rights | moral_rights | Waiver standard | N/A | N/A |
| Third-Party Payments | third_party_payments | Licensor responsible | Uncapped liability | Capped or shared |
| Audit Rights | audit_rights | Annual with notice | Not included | Included with favorable terms |

### ADMINISTRATIVE TERMS - Industry Standards
Note: If neither RED nor GREEN condition applies, use **YELLOW** (industry standard)

| Term | Field Name | Industry Standard (YELLOW) | RED if (worse) | GREEN if (better) |
|------|------------|----------------------------|----------------|-------------------|
| Accounting Frequency | accounting_frequency | Semi-annual | Annual or undefined | Quarterly |
| Objection Period | objection_period | 2-3 years | < 2 years | > 3 years |
| Litigation Deadline | litigation_deadline | 3-6 years | < 3 years | > 6 years |
| Assignment | assignment_rights | Mutual or limited | Label only | Mutual consent |
| Governing Law | governing_law | Label's jurisdiction | Disadvantages licensor | Licensor's jurisdiction |

## STEP 4: IDENTIFY RED FLAGS (ALL 18 - COMPREHENSIVE CHECK)

Check for these 18 red flags. Only flag if FOUND (except RF03 which flags absence).

### CRITICAL (Always RED)

**RF01: Fraction/Formula Royalty**
- Search terms: "fraction", "numerator", "denominator", "applicable fraction"
- Trigger: royalty = your rate ÷ artist rate (formula-based calculation)
- Impact: effective rate becomes ~10% of headline rate
- Example: "2% of 20% artist rate" = 0.1% effective

**RF02: Blind External Reference**
- Search terms: "per Recording Agreement", "see Exhibit", "in accordance with", "as per agreement between"
- Trigger: referenced document not provided or exhibit is blank
- Impact: unknown deductions, no transparency on actual terms

**RF03: No Audit Rights (ABSENCE FLAG)**
- Search terms: "audit", "inspect", "examine books", "accountant", "review records"
- Trigger: NONE of these terms found anywhere in document
- Impact: cannot verify royalty calculations ever

**RF04: Unlimited Indemnity Withholding**
- Search terms: "retain any monies", "withhold", "potential liability", "reserve against"
- Trigger: withholding allowed with no cap AND no time limit
- Impact: payments can be frozen indefinitely

**RF05: Net Zero Advance**
- Calculate: stated advance minus ALL deductions mentioned
- Trigger: net amount ≤ $0 or £0
- Impact: no actual upfront payment received

**RF06: Double Recoupment Gate**
- Search terms: "after [X] recoups", "following recoupment", "once Artist has recouped"
- Trigger: 2 or more conditions must be met before payment
- Impact: success of your work doesn't guarantee payment

### HIGH (RED)

**RF07: Stacked Undefined Deductions**
- Search terms: "proportionate reductions", "territorial reductions", "diminutions", "customary deductions"
- Trigger: multiple reduction types mentioned without specific caps or percentages
- Impact: 2% headline can become < 0.5% effective

**RF08: Sync Share via Fraction**
- Search terms: sync/licensing income + "fraction", "applicable fraction", "pro-rata"
- Trigger: licensing/sync income uses same fraction calculation as royalties
- Impact: "50% of sync" becomes ~5% effective

**RF09: Audio-Visual Reduction**
- Search terms: "audio-visual", "video" combined with "50%", "reduced by", "one-half"
- Trigger: AV royalty explicitly reduced by additional percentage beyond base calculation
- Impact: royalty is halved AGAIN for music videos and visual content

### MEDIUM (YELLOW or RED depending on severity)

**RF10: Escalation Excluded**
- Search terms: "without regard to escalation", "excluding escalation", "no escalation"
- Trigger: escalation provisions explicitly denied (not just missing)
- Color: RED if explicitly excluded, YELLOW if simply not mentioned

**RF11: Short Objection Period**
- Find: time limit to object/dispute statements
- Trigger: < 1 year = RED, 1-2 years = YELLOW, > 2 years = acceptable

**RF12: High Payment Threshold**
- Find: minimum payment amount before payout
- Trigger: > $200/£150 = RED, $100-200 = YELLOW

**RF13: Unlimited Remix Rights**
- Search terms: "any and all mixes and/or remixes", "unlimited derivatives"
- Trigger: no cap specified on number of remixes/versions allowed
- Impact: unlimited exploitation of your work in derivative forms

**RF14: Pro-Rata Compilation**
- Search terms: "pro-rated", "pro-rata", numerator/denominator + "tracks", "compilation"
- Trigger: royalty divided by track count on compilations
- Impact: 2% becomes 2%/15 = 0.13% on a 15-track compilation

**RF15: Reversionary Rights Waived**
- Search terms: "Section 203", "reversionary", "reversion" + "waive", "terminate"
- Trigger: explicit waiver of rights to reclaim copyright
- Impact: lose US law protection to reclaim work after 35 years

**RF16: Video Recoupment Gate**
- Search terms: "video", "music video" + "recoup", "recoupable", "costs recovered"
- Trigger: no AV payment until video production costs are fully recouped
- Impact: video royalties may never be paid if video costs were high

**RF17: All Services Bundled**
- Search terms: "no additional compensation", "inclusive of", "mixer", "engineer", "vocalist"
- Trigger: extra services (mixing, engineering, vocals) included without separate payment
- Impact: additional skilled work provided at no extra charge

**RF18: Unknown Deadlines**
- Search terms: "prior to Company's deadline", "before [other party's]", "reasonable time"
- Trigger: your obligations tied to unknown/undefined external dates
- Impact: cannot plan work or know when deadlines actually are

## STEP 5: CALCULATE EFFECTIVE RATE

If Applicable Fraction detected:
effective_rate = headline_rate × (headline_rate ÷ artist_rate)
Assume artist_rate = 20% if not specified.

Example calculation (for reference only - use actual document values):
- If headline is 2% and artist rate is 20%
- Fraction = 2% ÷ 20% = 0.10
- After typical deductions (~50%): effective ~1%

## OUTPUT FORMAT

Extract ALL values from THIS document. Return this JSON structure:

{{
  "agreement_summary": {{
    "type": "sample_clearance | work_for_hire | points_deal | beat_lease | beat_sale | co_production",
    "licensor": "[EXTRACT from this document]",
    "licensee": "[EXTRACT from this document]",
    "artist": "[EXTRACT from this document or 'Not specified']",
    "track": "[EXTRACT from this document or 'Not specified']",
    "effective_date": "[EXTRACT date or 'Not specified']",
    "exclusivity": "Exclusive | Non-Exclusive",
    "status": "Fully Executed | Partially Executed | No Signature"
  }},

  "overall_score": {{
    "rating": "FAVORABLE | NEUTRAL | UNFAVORABLE | HIGHLY_UNFAVORABLE",
    "red_flags": [COUNT from your analysis],
    "yellow_flags": [COUNT from your analysis],
    "green_flags": [COUNT from your analysis],
    "summary": "[Your assessment based on THIS document]"
  }},

  "royalty_analysis": {{
    "structure_type": "DIRECT | APPLICABLE_FRACTION",
    "headline_rate": "[EXTRACT from this document]",
    "effective_rate": "[CALCULATE based on document terms]",
    "explanation": "[Explain calculation using document terms]"
  }},

  "terms": {{
    "financial": {{
      // ONLY include nominal_fee if explicitly mentioned in document - OMIT entirely if not present
      "nominal_fee": {{
        "value": "[EXTRACT exact amount - OMIT this field if not mentioned]",
        "clause": "[Clause number where found]",
        "color": "YELLOW",
        "assessment": "[Your analysis]",
        "industry_standard": "£1-100 (UK/EU standard)"
      }},
      "advance": {{
        "value": "[EXTRACT exact amount or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "£3,000-10,000"
      }},
      "recoupable_advance": {{
        "value": "[EXTRACT recoupment terms: Fully recoupable / Non-recoupable / Cross-collateralized / Capped at X]",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis - is recoupment limited or unlimited?]",
        "industry_standard": "Fully recoupable from royalties"
      }},
      "royalty_rate": {{
        "value": "[EXTRACT exact rate or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "3-5% direct"
      }},
      "royalty_base": {{
        "value": "[EXTRACT definition or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "Net PPD with specified deductions"
      }},
      "payment_threshold": {{
        "value": "[EXTRACT amount or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "£50-100"
      }},
      "sync_share": {{
        "value": "[EXTRACT percentage or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "15-25% direct"
      }},
      "escalation": {{
        "value": "[EXTRACT terms or 'Not specified']",
        "clause": "[Clause number where found]",
        "color": "GREEN | YELLOW | RED",
        "assessment": "[Your analysis citing document]",
        "industry_standard": "Sales-based increases"
      }}
    }},
    "rights": {{
      "duration": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Perpetuity typical"}},
      "territory": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Worldwide"}},
      "media_scope": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "All media"}},
      "remix_rights": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "3-5 or additional fee"}},
      "name_likeness": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Credit + marketing"}},
      "touring_visuals": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Separate fee"}}
    }},
    "credit": {{
      "credit_format": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Sample credit format"}},
      "credit_placement": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Packaging + metadata"}},
      "credit_remedy": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Prospective cure"}}
    }},
    "legal": {{
      "warranties": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Standard package"}},
      "indemnification": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Mutual + capped"}},
      "moral_rights": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Waiver standard"}},
      "third_party_payments": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Licensor responsibility"}},
      "audit_rights": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Annual with notice"}}
    }},
    "administrative": {{
      "accounting_frequency": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Semi-annual"}},
      "objection_period": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "2-3 years"}},
      "litigation_deadline": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "3-6 years"}},
      "assignment_rights": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Mutual or limited"}},
      "governing_law": {{"value": "[EXTRACT]", "clause": "[Clause]", "color": "[COLOR]", "assessment": "[Analysis]", "industry_standard": "Negotiable"}}
    }}
  }},

  "red_flags": [
    // For EACH red flag detected (RF01-RF18), include an entry:
    {{
      "id": "RF01 | RF02 | RF03 | ... | RF18",
      "name": "[Short name from list above]",
      "severity": "CRITICAL | HIGH | MEDIUM",
      "clause": "[Clause reference from THIS document]",
      "quote": "[Exact quote from document proving this red flag exists - REQUIRED]",
      "trigger": "[What pattern in the document triggered this flag]",
      "impact": "[Financial or legal impact to producer/licensor]",
      "recommendation": "[Specific negotiation point or fix]"
    }}
    // IMPORTANT: RF03 (No Audit Rights) should be flagged if NO audit-related terms found
    // All other flags should ONLY be included if the triggering pattern IS found
  ],

  "negotiation_priorities": [
    {{
      "priority": 1,
      "term": "[field_name]",
      "current": "[Value from THIS document]",
      "target": "[Recommended target]",
      "impact": "CRITICAL | HIGH | MEDIUM"
    }}
  ],

  "field_ratings": {{
    "green": ["[List field names rated GREEN]"],
    "yellow": ["[List field names rated YELLOW]"],
    "red": ["[List field names rated RED]"]
  }},

  "general_assessment": "[Comprehensive assessment based ONLY on THIS agreement text. Do not reference other agreements.]"
}}

## SCORING

Count colors to determine overall rating:
- RED flags >= 5: HIGHLY_UNFAVORABLE
- RED flags 3-4: UNFAVORABLE
- RED flags 1-2 with YELLOW majority: NEUTRAL
- GREEN majority, RED = 0: FAVORABLE

## CRITICAL INSTRUCTIONS

1. **ISOLATION**: Treat this request as completely independent. Do not reference any other agreement or example.

2. **EXTRACTION ONLY**: Every value must come directly from THIS document. If you cannot find a term, use "Not specified".

3. **CITE SOURCES**: Every extracted value must include the clause number where it was found.

4. **QUOTE EVIDENCE**: For red flags, include a brief quote from the document proving the issue exists.

5. **RED FLAG DETECTION (RF01-RF18)** - Check ALL 18 red flags systematically:

   **CRITICAL FLAGS (Always add if detected):**
   - **RF01**: Search for "fraction", "numerator", "denominator" - flag if royalty uses formula
   - **RF02**: Search for "per Recording Agreement", "see Exhibit" - flag if external doc referenced but not provided
   - **RF03**: Search for "audit", "inspect", "examine books" - flag if NONE found (absence flag)
   - **RF04**: Search for "retain any monies", "withhold" - flag if no cap or time limit
   - **RF05**: Calculate advance minus deductions - flag if net ≤ 0
   - **RF06**: Search for multiple "recoup" conditions - flag if 2+ gates before payment

   **HIGH FLAGS:**
   - **RF07**: Search for "proportionate reductions", "diminutions" - flag if multiple undefined deductions
   - **RF08**: Search for sync + "fraction" - flag if sync uses fraction calculation
   - **RF09**: Search for "audio-visual" + "50%" or "reduced" - flag if AV has additional reduction

   **MEDIUM FLAGS:**
   - **RF10**: Search for "excluding escalation" - RED if excluded, YELLOW if missing
   - **RF11**: Check objection period - RED if < 1 year, YELLOW if 1-2 years
   - **RF12**: Check payment threshold - RED if > £200
   - **RF13**: Search for "any and all remixes" - flag if unlimited
   - **RF14**: Search for "pro-rata" + "compilation" - flag if royalty divided by track count
   - **RF15**: Search for "Section 203", "reversionary" + "waive" - flag if rights waived
   - **RF16**: Search for "video" + "recoup" - flag if video royalties gated by video costs
   - **RF17**: Search for "no additional compensation" + services - flag if services bundled
   - **RF18**: Search for "Company's deadline", "reasonable time" - flag if deadlines undefined

6. **AUTOMATIC RED FLAGS** - These are ALWAYS RED regardless of other factors:
   - **NO AUDIT RIGHTS**: If audit_rights is "Not specified" or no audit provision exists → ALWAYS add RF03 to red_flags AND set audit_rights.color = "RED"
   - **NO ESCALATION**: If escalation is "explicitly excluded" → set escalation.color = "RED"

7. **VALIDATION**: Before returning, verify:
   - Every value has a clause reference OR is marked "Not specified"
   - No values invented or assumed
   - Agreement type matches extracted content
   - Parties match document header
   - RF03 is included in red_flags if no audit provision found
   - All 18 red flags have been checked (even if most are not triggered)

AGREEMENT TEXT:
{text}"""


def parse_agreement_with_ai(text: str, agreement_type: str, file_path: str = None,
                           extraction_metadata: dict = None) -> dict:
    """
    Use AI to extract structured information from the agreement text.

    For producer agreements, uses the comprehensive MusicAgreementAnalyzer with
    18-red-flag detection system. For other agreement types, uses a simpler prompt.

    Falls back to sending PDF directly to Claude if text extraction failed.
    Returns a dictionary with parsed fields.

    Args:
        text: Extracted text from the document
        agreement_type: Type of agreement (e.g., "producer agreement")
        file_path: Path to the original document file
        extraction_metadata: Dict with extraction quality info (method, quality_score, warnings)
    """
    import base64

    # Clear the settings cache to pick up new env vars
    from app.settings.settings import get_settings as _get_settings
    _get_settings.cache_clear()

    settings = get_settings()
    extraction_metadata = extraction_metadata or {}

    print(f"[DEBUG] Parsing agreement, text length: {len(text) if text else 0}")
    print(f"[DEBUG] API key configured: {bool(settings.anthropic_api_key)}")
    print(f"[DEBUG] File path provided: {file_path}")
    print(f"[DEBUG] Agreement type: {agreement_type}")
    print(f"[DEBUG] Extraction quality: {extraction_metadata.get('quality_score', 'N/A')}")

    if not settings.anthropic_api_key:
        print("[DEBUG] No Anthropic API key configured, skipping AI parsing")
        return {"text_preview": text[:1000] if text else None}

    # Check if we need to use PDF fallback (text extraction failed)
    use_pdf_fallback = False
    if not text or len(text.strip()) < 50:
        if file_path and file_path.lower().endswith('.pdf'):
            print(f"[DEBUG] Text too short ({len(text.strip()) if text else 0} chars), will send PDF directly to Claude")
            use_pdf_fallback = True
        else:
            print(f"[DEBUG] Text too short for parsing: {len(text.strip()) if text else 0} chars")
            return {"text_preview": text[:1000] if text else None, "parse_error": "Document text too short"}

    # Use the comprehensive MusicAgreementAnalyzer for producer and publishing agreements
    if agreement_type.lower() in ["producer agreement", "publishing"] and not use_pdf_fallback:
        # Map frontend agreement type to analyzer agreement type constant
        agreement_type_hint = "PRODUCER_AGREEMENT" if agreement_type.lower() == "producer agreement" else "PUBLISHING_DEAL"
        return _parse_with_music_analyzer(text, settings.anthropic_api_key, extraction_metadata,
                                          agreement_type_hint=agreement_type_hint)

    # For management agreements or PDF fallback, use the legacy approach
    return _parse_with_legacy_prompt(text, agreement_type, file_path, use_pdf_fallback, settings.anthropic_api_key,
                                      extraction_metadata=extraction_metadata)


def _parse_with_music_analyzer(text: str, api_key: str, extraction_metadata: dict = None,
                                agreement_type_hint: str = None) -> dict:
    """
    Parse producer and publishing agreements using the comprehensive MusicAgreementAnalyzer.
    Returns analysis with 18-red-flag detection and structured terms.

    Args:
        text: Agreement text to analyze
        api_key: Anthropic API key
        extraction_metadata: Dict with extraction quality info
        agreement_type_hint: Optional hint about the agreement type (e.g., "PUBLISHING_DEAL")
    """
    try:
        print("[DEBUG] Using MusicAgreementAnalyzer for music agreement")
        analyzer = MusicAgreementAnalyzer(api_key=api_key)

        # Prepare options with agreement type hint
        options = {}
        if agreement_type_hint:
            options["dealTypeHint"] = agreement_type_hint
            print(f"[DEBUG] Agreement type hint: {agreement_type_hint}")

        # Add agreement_type_hint to extraction_metadata for post-processing
        extraction_metadata = extraction_metadata or {}
        if agreement_type_hint:
            extraction_metadata["agreement_type_hint"] = agreement_type_hint

        result = analyzer.analyze(text, options=options if options else None, extraction_metadata=extraction_metadata)

        if result.get("error"):
            print(f"[ERROR] Analyzer returned error: {result.get('error')}")
            return result

        # Map the new analyzer output to include backward-compatible field_ratings
        result = _add_backward_compatible_fields(result)

        # Store extraction metadata in result for frontend access
        if extraction_metadata:
            result["extraction_metadata"] = extraction_metadata

        # Check if too many NOT_FOUND terms suggest poor extraction
        not_found_count = 0
        total_terms = 0
        for section_data in result.get("terms", {}).values():
            for term_data in section_data.values():
                total_terms += 1
                if isinstance(term_data, dict) and term_data.get("value") == "NOT_FOUND":
                    not_found_count += 1

        # If >50% of terms are NOT_FOUND and extraction quality is low, suggest re-extraction
        extraction_quality = extraction_metadata.get("quality_score", 100) if extraction_metadata else 100
        if total_terms > 0 and not_found_count > total_terms * 0.5 and extraction_quality < 70:
            print(
                f"[WARNING] High NOT_FOUND rate ({not_found_count}/{total_terms}) with low extraction quality "
                f"({extraction_quality}/100) - consider re-extraction with Vision API"
            )
            # Add to result for frontend display
            result["analysis_warning"] = {
                "message": "Many terms not found - extraction quality may be low",
                "not_found_count": not_found_count,
                "total_terms": total_terms,
                "extraction_quality": extraction_quality,
                "suggestion": "Try re-extracting with Vision API for better results"
            }

        print(f"[DEBUG] MusicAgreementAnalyzer complete: {len(result.get('red_flags', []))} red flags detected")
        return result

    except Exception as e:
        import traceback
        print(f"[ERROR] MusicAgreementAnalyzer failed: {e}")
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        return {
            "text_preview": text[:1000] if text else None,
            "parse_error": f"Analyzer error: {str(e)}",
            "full_text": text,
            "extraction_metadata": extraction_metadata
        }


def _add_backward_compatible_fields(result: dict) -> dict:
    """
    Add backward-compatible field_ratings to the analyzer result.
    Maps the new color-coded terms to green/yellow/red arrays.
    """
    field_ratings = {"green": [], "yellow": [], "red": []}

    terms = result.get("terms", {})
    for section_key, section_terms in terms.items():
        if not isinstance(section_terms, dict):
            continue
        for field_name, term_data in section_terms.items():
            if not isinstance(term_data, dict):
                continue
            color = term_data.get("color", "GRAY")
            if color == "GREEN":
                field_ratings["green"].append(field_name)
            elif color == "YELLOW":
                field_ratings["yellow"].append(field_name)
            elif color == "RED":
                field_ratings["red"].append(field_name)

    result["field_ratings"] = field_ratings

    # Also map overall_score for backward compatibility
    overall = result.get("overall_assessment", {})
    result["overall_score"] = {
        "rating": overall.get("rating", "NEUTRAL"),
        "red_flags": len(result.get("red_flags", [])),
        "yellow_flags": overall.get("yellow_count", 0),
        "green_flags": overall.get("green_count", 0),
        "summary": overall.get("summary", "")
    }

    # Map agreement_summary for backward compatibility
    agreement = result.get("agreement", {})
    parties = agreement.get("parties", {})
    result["agreement_summary"] = {
        "type": agreement.get("type", "PRODUCER_AGREEMENT"),
        "licensor": parties.get("party_a", {}).get("name", ""),
        "licensee": parties.get("party_b", {}).get("name", ""),
        "artist": parties.get("artist", ""),
        "track": parties.get("track_or_project", ""),
        "effective_date": agreement.get("effective_date", ""),
        "exclusivity": "Exclusive",  # Default
        "status": "Fully Executed"  # Default
    }

    return result


def _parse_with_legacy_prompt(text: str, agreement_type: str, file_path: str, use_pdf_fallback: bool, api_key: str,
                               extraction_metadata: dict = None) -> dict:
    """
    Parse agreements using the legacy prompt-based approach.
    Used for publishing/management agreements and PDF fallback.

    Args:
        text: Agreement text to parse
        agreement_type: Type of agreement
        file_path: Path to original file (for PDF fallback)
        use_pdf_fallback: Whether to send PDF directly to Claude
        api_key: Anthropic API key
        extraction_metadata: Dict with extraction quality info (method, quality_score, warnings)
    """
    import base64

    extraction_metadata = extraction_metadata or {}

    try:
        import anthropic

        print("[DEBUG] Creating Anthropic client...")
        client = anthropic.Anthropic(api_key=api_key)

        # Use different prompt based on agreement type
        if agreement_type.lower() == "producer agreement":
            print("[DEBUG] Using producer agreement prompt")
            prompt = get_producer_agreement_prompt(text)
        else:
            print(f"[DEBUG] Using publishing/general agreement prompt for type: {agreement_type}")
            prompt = f"""You are a senior music industry attorney with 25+ years of experience analyzing recording contracts, publishing deals, and management agreements. You have represented major artists and work with labels like Universal, Sony, and Warner.

Analyze this {agreement_type} with EXTREME precision and COMPLETENESS. Extract EVERY detail - specific percentages, dollar amounts, dates, party names, and ALL terms as written in the document.

CRITICAL: NEVER omit or summarize information. If the document mentions multiple items (e.g., multiple deductions, multiple approval requirements), list ALL of them. Be exhaustive.

Return ONLY a valid JSON object with these fields:

{{
    "agreement_type": "{agreement_type}",
    "subtype": "Identify the specific deal structure. For publishing: 'Administration' (no ownership transfer), 'Co-publishing' (shared ownership 50/50), 'Full Publishing' (full transfer), 'Hybrid' (administration with some assignment of rights/ownership), 'Sub-publishing'. For recording: Exclusive Artist, License, Distribution, Joint Venture. For producer: Work for Hire, Points Deal, Flat Fee. For management: Full Service, Consulting, 360. Use 'Hybrid' when you see BOTH admin-style terms AND assignment/transfer of rights or ownership.",
    "exclusivity": "Exactly 'Exclusive' or 'Non-Exclusive'. Look for language like 'sole and exclusive', 'first right of refusal', or 'non-exclusive license'.",
    "status": "Check signature blocks at end of document. 'Fully Executed' if all parties signed with dates, 'Partially Executed' if some signatures present, 'No Signature' if signature lines are blank or document is a draft.",
    "assigner": "The EXACT legal name of the party granting rights. This is typically the artist, songwriter, producer, or rights holder. Include 'p/k/a' or stage names if mentioned.",
    "assignee": "The EXACT legal name of the party receiving/acquiring rights. This is typically the label, publisher, manager, or distributor. Include parent company if mentioned.",
    "royalty_rates": {{
        "mechanical_domestic": "ONLY percentage and BASE TYPE. DO NOT say 'mechanical', 'domestic', or 'international'. BASE must be: 'Net Receipts', 'Gross Receipts', or 'at source'. Examples: '85% of Net Receipts', '80% at source', '75% of Gross'. If agreement has one blanket rate, use that. Use null ONLY if no royalty rate found.",
        "mechanical_international": "ONLY percentage and BASE TYPE. DO NOT say 'domestic' or 'international'. Use same base format as domestic. If no domestic/international differentiation, use SAME value as domestic. Use null ONLY if domestic is null.",
        "performance_domestic": "ONLY percentage and BASE TYPE. CRITICAL: If the agreement has a blanket royalty rate (e.g., '85% of Net Receipts') without separate mechanical/performance rates, USE THE SAME RATE for performance as mechanical. Do NOT say '0%' or 'paid direct' - that refers to PRO payments which are separate. The performance rate here is what the PUBLISHER pays the writer from publisher-collected performance income. If no separate rate specified, COPY the mechanical_domestic value exactly.",
        "performance_international": "ONLY percentage and BASE TYPE. If no domestic/international split, use SAME as performance_domestic. Use null ONLY if performance_domestic is null.",
        "sync": "ONLY percentage and BASE TYPE. DO NOT say 'sync', 'synchronization'. Use same base format. If not specified separately, use the blanket royalty rate. Use null ONLY if no rates found.",
        "master": "FOR RECORDING ONLY: ONLY percentage and BASE TYPE. DO NOT say 'master'. Use null for publishing agreements.",
        "other": "Other income with percentage and BASE TYPE. Examples: 'Print: 10% of Net'. Use null if none specified."
    }},
    "costs": {{
        "advance": "Advance amount and recoupment terms (e.g., '$10,000 recoupable from writer share'). Use null if no advance.",
        "marketing": "FOR RECORDING: Marketing/promotion costs (e.g., 'Up to $5,000 recoupable'). Use null for publishing agreements.",
        "distribution_fee": "FOR RECORDING: Distribution fee percentage (e.g., '15% off the top'). Use null for publishing agreements.",
        "recording_costs": "FOR RECORDING: Recording budget terms (e.g., '$50,000 budget, 100% recoupable'). Use null for publishing agreements.",
        "other_deductions": "LIST ALL deductible costs mentioned - be EXHAUSTIVE. Include EVERY item: demo costs, copyright registration fees, legal fees, sample clearance, marketing contributions, third-party payments, co-writer shares, producer fees, etc. Format each with recoupability status. Example: 'Demo costs: recoupable, Copyright registration: non-recoupable, Sample clearance: recoupable at 50%, Third-party payments: deducted off the top'. NEVER omit any deduction mentioned in the document. Use null ONLY if absolutely no deductions are mentioned."
    }},
    "audit_rights": {{
        "statement_frequency": "How often statements are provided (e.g., 'Quarterly, within 45 days of quarter end'). Use null if not specified.",
        "audit_window": "Time period to request audit after receiving statement (e.g., '2 years from statement date'). Use null if not specified.",
        "objection_period": "Deadline to dispute/object to statement (e.g., '90 days or statement becomes final and binding'). Use null if not specified.",
        "audit_frequency": "How often audits are permitted (e.g., 'Once per calendar year with 30 days written notice'). Use null if not specified.",
        "audit_costs": "Who pays for audit (e.g., 'Artist pays unless discrepancy exceeds 10%, then Company pays'). Use null if not specified."
    }},
    "approvals": "List as comma-separated string. Example: 'Sync licenses over $5,000, major label licenses, film/TV placements requiring writer consent'.",
    "term": "State as a single string including termination rights. Example: 'Initial term: 2 years. Two 1-year options at Publisher discretion. Termination: Writer may terminate with 90 days written notice if Publisher fails to exploit works for 12 months. Retention period: 10 years post-term for works created during term.'",
    "renewal": {{
        "auto_renews": "Does the agreement auto-renew? 'Yes' or 'No'. Look for 'automatically renew', 'successive terms', 'evergreen'. Use null if unclear.",
        "renewal_period": "How long is each renewal period? (e.g., '1 year', '12 months', 'successive 1-year terms'). Use null if not auto-renewing.",
        "termination_notice_days": "How many days notice required to terminate/prevent renewal? Extract as NUMBER ONLY (e.g., '30', '60', '90'). Use null if not specified.",
        "next_renewal_date": "When is the next renewal/extension date? Format as YYYY-MM-DD if determinable. Use null if cannot calculate."
    }},
    "collection_period": "Also called 'retention period'. If NO collection/retention period exists, output exactly 'No collection period'. If one exists, describe duration (e.g., '6 months post-term', '2 years after expiration'). Example: 'Publisher retains collection rights for 5 years post-term.'",
    "effective_date": "The contract commencement date. Look for 'Effective Date', 'as of', or date next to signatures. Format as YYYY-MM-DD if possible.",
    "expiration_date": "Calculate from term if determinable. Format as YYYY-MM-DD. Use 'Life of Copyright' or 'Perpetual' if applicable.",
    "territory": "Exact territorial scope: 'World', 'Universe', specific countries listed, or exclusions (e.g., 'Worldwide excluding Japan').",
    "general_assessment": "As an experienced music attorney, provide a BALANCED 4-6 sentence assessment. PRIORITY ORDER for weighting terms (most to least important): (1) TERM LENGTH & TERMINATION FLEXIBILITY - short terms with easy exit are VERY valuable, often offsetting other concerns. (2) ROYALTY RATES - 80%+ is standard, below 50% is concerning. (3) Collection period and rights retention. (4) Audit rights (least important - rarely exercised in practice). RECOGNIZE TRADE-OFFS: A deal with permanent 20% publisher share BUT short 2-month term with 30-day termination is a BALANCED trade-off, not predatory. The flexibility to exit quickly compensates for the retention. Use measured language. RATINGS: Artist-Favorable (terms clearly favor artist), Balanced (fair trade-offs on both sides), Company-Favorable (terms lean toward company but not exploitative). Reserve 'Predatory' ONLY for truly egregious cases with MULTIPLE intentional traps: life-of-copyright terms WITH no termination AND sub-50% royalties AND hidden fees.",
    "field_ratings": {{
        "green": [],
        "yellow": [],
        "red": []
    }}
}}

CRITICAL INSTRUCTION FOR field_ratings (YOU MUST POPULATE THESE ARRAYS):

RATE EVERY APPLICABLE FIELD. Valid field names:
- Royalty fields: mechanical_domestic, mechanical_international, performance_domestic, performance_international, sync, master, other
- Rights fields: subtype, exclusivity, approvals
- Term fields: term, collection_period, expiration_date
- Audit fields: audit_window, objection_period, audit_costs, audit_frequency, statement_frequency
- Financial: advance

GREEN (artist-favorable) - Add ALL fields that favor the artist:
- ANY royalty rate 80%+ (add mechanical_domestic, mechanical_international, performance_domestic, performance_international, sync as applicable)
- term: Initial term under 1 year = ALWAYS GREEN (e.g., "2 months", "6 months", "month-to-month"). Recoupment-based extensions are standard/acceptable, NOT a trap. Short renewable terms with easy termination notice (30-60 days) are artist-favorable.
- collection_period: No collection/retention period = ALWAYS GREEN. "No collection period" = GREEN.
- audit_window: 2+ years to request audit
- audit_costs: Company pays if discrepancy found (artist-favorable split)
- Easy termination with 30-60 days notice
- Approval rights for sync/licensing

TERM RATING EXAMPLES:
- "2 months initial, renewable in 2-month periods, 30 days notice" → term = GREEN (very short, easy exit)
- "6 months with option to renew" → term = GREEN
- "1 year with 60 days termination notice" → term = GREEN
- "3 years firm" → term = YELLOW
- "Life of copyright" → term = RED

YELLOW (concerning) - Add ALL fields with below-standard terms:
- ANY royalty rate 50-79% (rate EACH royalty field separately)
- Terms 3-5 years
- collection_period: 12 months or less = YELLOW (e.g., "6 months post-term", "1 year after expiration")
- audit_window: exactly 2 years (borderline)
- audit_costs: Mixed terms (e.g., "Artist pays but company reimburses if discrepancy exceeds 15%")

RED (severe red flags) - CRITICAL: Add ALL fields with predatory terms:
- ANY royalty rate under 50%
- audit_window: 1 year or less, OR not specified = ALWAYS RED (artist loses right to audit old statements quickly)
- audit_costs: Owner/artist pays audit costs in ALL scenarios = ALWAYS RED. Industry standard is company pays if underpaid by 10%+. If there's NO provision for company to pay when underpayment found, it's RED.
- Life of copyright with no termination
- collection_period: Over 12 months = RED (e.g., "2 years", "5 years post-term", "life of copyright collection rights")
- No approval rights

COLLECTION PERIOD RATING EXAMPLES:
- "No collection period" or no retention clause → collection_period = GREEN
- "6 months post-term" → collection_period = YELLOW
- "12 months after Term end" → collection_period = YELLOW (exactly 12 months = YELLOW)
- "1 year after expiration" → collection_period = YELLOW (1 year = 12 months = YELLOW)
- "18 months post-term" → collection_period = RED (over 12 months)
- "2 years post-term" → collection_period = RED
- "5 years retention" → collection_period = RED
- "Life of copyright collection rights" → collection_period = RED

AUDIT WINDOW RATING EXAMPLES (CRITICAL - 1 year or less in ANY phrasing = RED):
- "1 year from statement" → audit_window = RED
- "12 months" → audit_window = RED
- "one year" → audit_window = RED
- "within 1 year" → audit_window = RED
- "1 year after receipt" → audit_window = RED
- "within twelve months" → audit_window = RED
- "1 year period" → audit_window = RED
- "not specified" or "silent" → audit_window = RED
- "2 years" → audit_window = YELLOW (borderline)
- "3 years" → audit_window = GREEN

AUDIT COSTS RATING EXAMPLES:
- "Owner bears all audit costs" → audit_costs = RED
- "Artist responsible for audit expenses" → audit_costs = RED
- "Owner pays, subject to approval for auditors" → audit_costs = RED (owner still pays ALL costs)
- "Owner pays unless discrepancy found" → audit_costs = RED (no threshold = still unfavorable)
- "Company pays if discrepancy exceeds 10%" → audit_costs = GREEN
- "Company pays if underpayment exceeds 5%" → audit_costs = GREEN

IMPORTANT: Rate EACH royalty field individually. If mechanical_domestic is 85% (green) but sync is 50% (yellow), add mechanical_domestic to green AND sync to yellow.

Example: {{"green": ["term", "mechanical_domestic", "mechanical_international", "performance_domestic", "performance_international", "audit_window"], "yellow": ["sync"], "red": ["collection_period"]}}

IMPORTANT REMINDERS:
1. Be EXHAUSTIVE - never omit information. List ALL deductions, ALL approval requirements, ALL terms mentioned.
2. Use exact quotes and figures from the document.
3. Rate EVERY applicable field in field_ratings - especially ALL royalty fields (domestic AND international).
4. If a field truly cannot be determined from the text, use null.

AGREEMENT TEXT:
{text}"""

        print("[DEBUG] Calling Anthropic API...")

        # Build message content based on whether we're using PDF fallback
        if use_pdf_fallback:
            # Read and encode the PDF file
            with open(file_path, "rb") as pdf_file:
                pdf_data = base64.standard_b64encode(pdf_file.read()).decode("utf-8")

            print(f"[DEBUG] Sending PDF directly to Claude (size: {len(pdf_data)} bytes base64)")

            # Create prompt without the text (since we're sending the PDF)
            pdf_prompt = prompt.replace(f"\n\nAGREEMENT TEXT:\n{text}", "")
            pdf_prompt += "\n\nAnalyze the attached PDF document."

            message_content = [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_data
                    }
                },
                {
                    "type": "text",
                    "text": pdf_prompt
                }
            ]

            message = client.messages.create(
                model="claude-sonnet-4-20250514",  # Use Sonnet for PDF vision capability
                max_tokens=4000,
                messages=[{"role": "user", "content": message_content}],
            )
        else:
            # Add JSON-only instruction
            json_prompt = "CRITICAL: Output ONLY valid JSON. No preamble, no explanation, no markdown. Start with { and end with }.\n\n" + prompt

            # Use Sonnet for producer agreements (complex extraction requires better model)
            if agreement_type.lower() == "producer agreement":
                print("[DEBUG] Using Sonnet for producer agreement (complex extraction)")
                message = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=16000,
                    messages=[{"role": "user", "content": json_prompt}],
                )
            else:
                message = client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=8000,
                    messages=[{"role": "user", "content": json_prompt}],
                )

        # Extract JSON from response
        response_text = message.content[0].text.strip()
        print(f"[DEBUG] Got response, length: {len(response_text)}")
        # Debug: Print first 2000 chars of response to see structure
        print(f"[DEBUG] Response preview: {response_text[:2000]}")

        # Try to parse the JSON
        # Handle case where response might have markdown code blocks
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
        else:
            # Find JSON object by locating first { and matching closing }
            first_brace = response_text.find('{')
            if first_brace != -1:
                # Count braces to find the matching closing brace
                brace_count = 0
                in_string = False
                escape_next = False
                end_pos = first_brace

                for i, char in enumerate(response_text[first_brace:], start=first_brace):
                    if escape_next:
                        escape_next = False
                        continue
                    if char == '\\':
                        escape_next = True
                        continue
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                    if not in_string:
                        if char == '{':
                            brace_count += 1
                        elif char == '}':
                            brace_count -= 1
                            if brace_count == 0:
                                end_pos = i
                                break

                if brace_count == 0 and end_pos > first_brace:
                    response_text = response_text[first_brace:end_pos + 1]
                    print(f"[DEBUG] Extracted JSON from position {first_brace} to {end_pos}")

        # Try to parse JSON, with repair attempts if it fails
        try:
            parsed = json.loads(response_text)
        except json.JSONDecodeError as initial_error:
            print(f"[DEBUG] Initial JSON parse failed: {initial_error}, attempting repair...")

            # Attempt 1: Fix truncated strings by finding the last complete object/array
            repaired = response_text

            # Remove trailing incomplete strings (after last complete field)
            # Find the last properly closed structure
            import re

            # Try to find and fix common issues
            # Issue: Truncated string values - close them
            # Look for patterns like "value": "some text without closing quote
            lines = repaired.split('\n')
            fixed_lines = []
            for i, line in enumerate(lines):
                # Check for unclosed string at end of line (not followed by comma/brace)
                if re.search(r':\s*"[^"]*$', line) and i < len(lines) - 1:
                    # Line has unclosed string, check if next line continues it
                    next_line = lines[i + 1] if i + 1 < len(lines) else ""
                    if not next_line.strip().startswith('"'):
                        # Close the string and add comma
                        line = line + '",'
                fixed_lines.append(line)
            repaired = '\n'.join(fixed_lines)

            # Attempt 2: Ensure proper closing braces/brackets
            open_braces = repaired.count('{') - repaired.count('}')
            open_brackets = repaired.count('[') - repaired.count(']')

            # Remove trailing comma before closing
            repaired = re.sub(r',\s*$', '', repaired.rstrip())

            # Add missing closing brackets/braces
            if open_brackets > 0:
                repaired += ']' * open_brackets
            if open_braces > 0:
                repaired += '}' * open_braces

            try:
                parsed = json.loads(repaired)
                print(f"[DEBUG] JSON repair successful!")
            except json.JSONDecodeError as repair_error:
                print(f"[DEBUG] JSON repair failed: {repair_error}")
                # Last resort: try to extract valid JSON objects
                raise initial_error

        # Debug: Log the top-level keys and terms structure
        print(f"[DEBUG] Parsed JSON top-level keys: {list(parsed.keys())}")
        if "terms" in parsed:
            terms = parsed.get("terms", {})
            print(f"[DEBUG] terms sub-keys: {list(terms.keys()) if isinstance(terms, dict) else 'NOT A DICT'}")
            for section in ["financial", "rights", "credit", "legal", "administrative"]:
                section_data = terms.get(section, {}) if isinstance(terms, dict) else {}
                field_count = len(section_data) if isinstance(section_data, dict) else 0
                print(f"[DEBUG] terms.{section}: {field_count} fields")
        else:
            print("[DEBUG] WARNING: No 'terms' key in parsed response!")

        # Clean up field_ratings - remove any placeholder text the AI left in
        if "field_ratings" in parsed and isinstance(parsed["field_ratings"], dict):
            valid_fields = {
                # Publishing/general fields
                "subtype", "exclusivity", "mechanical_domestic", "mechanical_international",
                "performance_domestic", "performance_international", "sync", "master", "other", "advance",
                "term", "collection_period", "audit_window", "objection_period",
                # Producer agreement fields
                "base_rate", "producer_fee", "total_advance", "recoupable_amount", "all_in_fund",
                "backend_percentage", "master_ownership", "work_for_hire", "composition_rights",
                "reversion_rights", "letter_of_direction", "audit_costs", "credit_format",
                "royalty_structure", "escalations",
                          "audit_costs", "audit_frequency", "statement_frequency", "approvals", "expiration_date"}

            # Clean green array
            if "green" in parsed["field_ratings"]:
                parsed["field_ratings"]["green"] = [
                    f for f in parsed["field_ratings"]["green"]
                    if isinstance(f, str) and f in valid_fields
                ]
            else:
                parsed["field_ratings"]["green"] = []

            # Clean yellow array
            if "yellow" in parsed["field_ratings"]:
                parsed["field_ratings"]["yellow"] = [
                    f for f in parsed["field_ratings"]["yellow"]
                    if isinstance(f, str) and f in valid_fields
                ]
            else:
                parsed["field_ratings"]["yellow"] = []

            # Clean red array
            if "red" in parsed["field_ratings"]:
                parsed["field_ratings"]["red"] = [
                    f for f in parsed["field_ratings"]["red"]
                    if isinstance(f, str) and f in valid_fields
                ]
            else:
                parsed["field_ratings"]["red"] = []

            print(f"[DEBUG] Field ratings - green: {parsed['field_ratings']['green']}, yellow: {parsed['field_ratings']['yellow']}, red: {parsed['field_ratings']['red']}")

            # POST-PROCESSING: Enforce audit_window rating based on value
            # This reduces variability from AI interpretation
            import re
            audit_window_val = parsed.get("audit_rights", {}).get("audit_window", "") if isinstance(parsed.get("audit_rights"), dict) else ""
            print(f"[DEBUG] Post-process audit_window raw value: '{audit_window_val}'")
            if audit_window_val:
                audit_window_lower = str(audit_window_val).lower()

                # Remove audit_window from all rating arrays first
                for rating_type in ["green", "yellow", "red"]:
                    if "audit_window" in parsed["field_ratings"][rating_type]:
                        parsed["field_ratings"][rating_type].remove("audit_window")

                # Use regex for more robust matching
                # Check for 3+ year patterns first - should be GREEN
                three_plus_match = re.search(r'\b([3-9]|[1-9]\d+)\s*[-]?\s*years?\b', audit_window_lower) or \
                                   re.search(r'\b(three|four|five|six|seven|eight|nine|ten)\s*years?\b', audit_window_lower) or \
                                   re.search(r'\b(36|48|60)\s*months?\b', audit_window_lower)

                # Check for 2 year patterns - should be YELLOW
                two_year_match = re.search(r'\b2\s*[-]?\s*years?\b', audit_window_lower) or \
                                 re.search(r'\btwo\s*years?\b', audit_window_lower) or \
                                 re.search(r'\b24\s*months?\b', audit_window_lower)

                # Check for 1 year or less patterns - should be RED
                # Match: "1 year", "one year", "12 months", "(1) year", "1-year", "one (1) year", etc.
                one_year_match = re.search(r'\b1\s*[-]?\s*years?\b', audit_window_lower) or \
                                 re.search(r'\(\s*1\s*\)\s*years?', audit_window_lower) or \
                                 re.search(r'\bone\s*(\(\s*1\s*\))?\s*years?\b', audit_window_lower) or \
                                 re.search(r'\b12\s*[-]?\s*months?\b', audit_window_lower) or \
                                 re.search(r'\btwelve\s*months?\b', audit_window_lower) or \
                                 re.search(r'\b365\s*days?\b', audit_window_lower)

                print(f"[DEBUG] Pattern matches - 3+yr: {bool(three_plus_match)}, 2yr: {bool(two_year_match)}, 1yr: {bool(one_year_match)}")

                # Apply correct rating based on pattern (check in order: 3+, 2, 1, unspecified)
                if three_plus_match and not one_year_match:
                    parsed["field_ratings"]["green"].append("audit_window")
                    print(f"[DEBUG] Post-process: audit_window -> GREEN (3+ years)")
                elif two_year_match and not one_year_match:
                    parsed["field_ratings"]["yellow"].append("audit_window")
                    print(f"[DEBUG] Post-process: audit_window -> YELLOW (2 years)")
                elif one_year_match:
                    parsed["field_ratings"]["red"].append("audit_window")
                    print(f"[DEBUG] Post-process: audit_window -> RED (1 year or less)")
                elif "not specified" in audit_window_lower or audit_window_lower == "null":
                    parsed["field_ratings"]["red"].append("audit_window")
                    print(f"[DEBUG] Post-process: audit_window not specified -> RED")
                else:
                    # Default to RED if we can't parse it (safer assumption)
                    parsed["field_ratings"]["red"].append("audit_window")
                    print(f"[DEBUG] Post-process: audit_window unparseable -> RED (defaulting)")

            print("[DEBUG] Post-process: After audit_window block, continuing to performance rates...")

            # POST-PROCESSING: Fix performance rates if AI returned "0%" or "paid direct"
            # When no separate rate is specified, performance should match mechanical
            royalty_rates = parsed.get("royalty_rates", {})
            if isinstance(royalty_rates, dict):
                mech_dom = royalty_rates.get("mechanical_domestic", "")
                perf_dom = royalty_rates.get("performance_domestic", "")

                # Check if performance_domestic looks wrong (0%, paid direct, etc.)
                if perf_dom and isinstance(perf_dom, str):
                    perf_lower = perf_dom.lower()
                    # Detect incorrect values: 0%, paid direct, writer's share, etc.
                    if ("0%" in perf_lower or
                        "paid direct" in perf_lower or
                        "writer's share" in perf_lower or
                        "100% writer" in perf_lower or
                        "directly to writer" in perf_lower or
                        "n/a" in perf_lower):
                        # Copy from mechanical if mechanical has a real value
                        if mech_dom and isinstance(mech_dom, str) and "%" in mech_dom and "0%" not in mech_dom.lower():
                            print(f"[DEBUG] Post-process: Fixing performance_domestic from '{perf_dom}' to '{mech_dom}' (copying mechanical)")
                            parsed["royalty_rates"]["performance_domestic"] = mech_dom
                            # Also fix international
                            mech_intl = royalty_rates.get("mechanical_international", "")
                            if mech_intl and isinstance(mech_intl, str) and "%" in mech_intl:
                                perf_intl = royalty_rates.get("performance_international", "")
                                if perf_intl and isinstance(perf_intl, str):
                                    perf_intl_lower = perf_intl.lower()
                                    if ("0%" in perf_intl_lower or "paid direct" in perf_intl_lower or
                                        "writer's share" in perf_intl_lower or "n/a" in perf_intl_lower):
                                        print(f"[DEBUG] Post-process: Fixing performance_international from '{perf_intl}' to '{mech_intl}'")
                                        parsed["royalty_rates"]["performance_international"] = mech_intl

                            # Update field_ratings to rate performance same as mechanical
                            # Check if mechanical_domestic is in any rating array
                            for rating_type in ["green", "yellow", "red"]:
                                if "mechanical_domestic" in parsed["field_ratings"][rating_type]:
                                    # Add performance_domestic to same rating if not already there
                                    if "performance_domestic" not in parsed["field_ratings"][rating_type]:
                                        # First remove from other arrays
                                        for other_rating in ["green", "yellow", "red"]:
                                            if "performance_domestic" in parsed["field_ratings"][other_rating]:
                                                parsed["field_ratings"][other_rating].remove("performance_domestic")
                                        parsed["field_ratings"][rating_type].append("performance_domestic")
                                        print(f"[DEBUG] Post-process: performance_domestic rating -> {rating_type} (same as mechanical)")
                                    break

            # POST-PROCESSING: Enforce audit_costs rating based on value
            audit_costs_val = parsed.get("audit_rights", {}).get("audit_costs", "") if isinstance(parsed.get("audit_rights"), dict) else ""
            if audit_costs_val:
                audit_costs_lower = str(audit_costs_val).lower()
                # Remove from all arrays first
                for rating_type in ["green", "yellow", "red"]:
                    if "audit_costs" in parsed["field_ratings"][rating_type]:
                        parsed["field_ratings"][rating_type].remove("audit_costs")

                # Check if company pays when discrepancy found (GREEN)
                company_pays_patterns = ["company pays if", "administrator pays if", "publisher pays if",
                                         "label pays if", "discrepancy exceeds", "underpayment exceeds"]
                # Check if owner/artist pays all costs (RED)
                owner_pays_patterns = ["owner pays", "owner bears", "artist pays", "artist bears",
                                      "writer pays", "writer bears", "owner responsible", "artist responsible",
                                      "at owner", "at artist", "owner's expense", "artist's expense"]

                if any(pattern in audit_costs_lower for pattern in company_pays_patterns):
                    parsed["field_ratings"]["green"].append("audit_costs")
                    print(f"[DEBUG] Post-process: audit_costs -> GREEN (company pays if discrepancy)")
                elif any(pattern in audit_costs_lower for pattern in owner_pays_patterns):
                    parsed["field_ratings"]["red"].append("audit_costs")
                    print(f"[DEBUG] Post-process: audit_costs '{audit_costs_val}' -> RED (owner pays)")

            # POST-PROCESSING: Enforce collection_period rating based on value
            # Shorter collection periods are better for the artist
            collection_period_val = parsed.get("collection_period", "")
            if collection_period_val and isinstance(collection_period_val, str):
                collection_lower = collection_period_val.lower()
                print(f"[DEBUG] Post-process collection_period raw value: '{collection_period_val}'")

                # Remove from all arrays first
                for rating_type in ["green", "yellow", "red"]:
                    if "collection_period" in parsed["field_ratings"][rating_type]:
                        parsed["field_ratings"][rating_type].remove("collection_period")

                # Check for perpetual/life of copyright (RED)
                perpetual_patterns = ["perpetual", "perpetuity", "life of copyright", "forever", "in perpetuity", "indefinite"]
                is_perpetual = any(pattern in collection_lower for pattern in perpetual_patterns)

                # Check for 3+ years (RED)
                three_plus_years = re.search(r'\b([3-9]|[1-9]\d+)\s*[-]?\s*years?\b', collection_lower) or \
                                   re.search(r'\b(three|four|five|six|seven|eight|nine|ten)\s*years?\b', collection_lower) or \
                                   re.search(r'\b(36|48|60)\s*months?\b', collection_lower)

                # Check for 2 year patterns (YELLOW)
                two_year_match = re.search(r'\b2\s*[-]?\s*years?\b', collection_lower) or \
                                 re.search(r'\btwo\s*years?\b', collection_lower) or \
                                 re.search(r'\b24\s*months?\b', collection_lower)

                # Check for 18 months (YELLOW)
                eighteen_months = re.search(r'\b18\s*months?\b', collection_lower) or \
                                  re.search(r'\beighteen\s*months?\b', collection_lower)

                # Check for 1 year or 12 months (GREEN - good for artist)
                one_year_or_less = re.search(r'\b1\s*[-]?\s*years?\b', collection_lower) or \
                                   re.search(r'\bone\s*(\(\s*1\s*\))?\s*years?\b', collection_lower) or \
                                   re.search(r'\b12\s*[-]?\s*months?\b', collection_lower) or \
                                   re.search(r'\btwelve\s*months?\b', collection_lower) or \
                                   re.search(r'\b([1-9]|1[01])\s*months?\b', collection_lower)

                # Apply rating based on WORST case when multiple values present
                # (e.g., "12 months (US), 18 months (International)" should be YELLOW due to 18 months)
                if is_perpetual:
                    parsed["field_ratings"]["red"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period -> RED (perpetual)")
                elif three_plus_years:
                    parsed["field_ratings"]["red"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period -> RED (3+ years)")
                elif two_year_match:
                    parsed["field_ratings"]["yellow"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period -> YELLOW (2 years)")
                elif eighteen_months:
                    parsed["field_ratings"]["yellow"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period -> YELLOW (18 months)")
                elif one_year_or_less:
                    parsed["field_ratings"]["green"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period -> GREEN (12 months or less)")
                else:
                    # Default to yellow if we can't parse
                    parsed["field_ratings"]["yellow"].append("collection_period")
                    print(f"[DEBUG] Post-process: collection_period unparseable -> YELLOW (defaulting)")

            print("[DEBUG] Post-process: Entering RF03 detection block...")

            # POST-PROCESSING: Enforce RF03 (No Audit Rights) for producer agreements
            # Check both the AI-extracted value AND the original text for audit provisions
            has_no_audit = False
            audit_value = ""

            # Method 1: Check AI-extracted terms.legal.audit_rights
            if "terms" in parsed and isinstance(parsed.get("terms"), dict):
                legal_terms = parsed["terms"].get("legal", {})
                if isinstance(legal_terms, dict):
                    audit_rights_term = legal_terms.get("audit_rights", {})
                    if isinstance(audit_rights_term, dict):
                        audit_value = str(audit_rights_term.get("value", "")).lower()
                    elif isinstance(audit_rights_term, str):
                        audit_value = audit_rights_term.lower()

                    # Check if audit rights are missing/not specified from AI extraction
                    no_audit_patterns = ["not specified", "not included", "not mentioned", "none", "n/a"]
                    has_no_audit = any(pattern in audit_value for pattern in no_audit_patterns) or audit_value.strip() == ""
                    print(f"[DEBUG] Post-process RF03: AI audit_rights value = '{audit_value}', has_no_audit = {has_no_audit}")

            # Method 2: Also search the original document text for audit-related terms
            # This catches cases where AI failed to extract but audit terms exist
            full_text = parsed.get("full_text", "") or text or ""
            text_lower = full_text.lower() if full_text else ""

            audit_keywords = ["audit", "inspect", "examine books", "examine records", "accountant", "review records", "accounting records"]
            text_has_audit_terms = any(keyword in text_lower for keyword in audit_keywords)
            print(f"[DEBUG] Post-process RF03: text_has_audit_terms = {text_has_audit_terms}")

            # If text DOES have audit terms but AI said "not specified", trust the text
            # If text DOES NOT have audit terms, flag RF03 regardless of AI extraction
            if not text_has_audit_terms:
                has_no_audit = True
                print(f"[DEBUG] Post-process RF03: No audit keywords found in document text - flagging RF03")
            elif has_no_audit and text_has_audit_terms:
                # AI missed it but text has it - don't flag
                has_no_audit = False
                print(f"[DEBUG] Post-process RF03: AI missed audit rights but text has audit terms - NOT flagging")

            if has_no_audit:
                print(f"[DEBUG] Post-process: NO AUDIT RIGHTS DETECTED - adding RF03")

                # Ensure red_flags array exists
                if "red_flags" not in parsed:
                    parsed["red_flags"] = []
                if not isinstance(parsed["red_flags"], list):
                    parsed["red_flags"] = []

                # Check if RF03 already exists
                rf03_exists = any(
                    isinstance(rf, dict) and rf.get("id") == "RF03"
                    for rf in parsed["red_flags"]
                )

                if not rf03_exists:
                    parsed["red_flags"].append({
                        "id": "RF03",
                        "name": "No Audit Rights",
                        "severity": "CRITICAL",
                        "clause": "N/A - Not found in document",
                        "quote": "No audit provision found in agreement",
                        "trigger": "Absence of audit, inspect, examine books, or review records language",
                        "impact": "Cannot verify royalty calculations - no way to confirm you are being paid correctly",
                        "recommendation": "Add audit rights clause: right to audit books once per year with 30 days notice"
                    })
                    print(f"[DEBUG] Post-process: Added RF03 (No Audit Rights) to red_flags")

                # Also ensure audit_rights is marked as RED in field_ratings
                if "audit_rights" not in parsed["field_ratings"]["red"]:
                    # Remove from other arrays first
                    for rating_type in ["green", "yellow"]:
                        if "audit_rights" in parsed["field_ratings"][rating_type]:
                            parsed["field_ratings"][rating_type].remove("audit_rights")
                    parsed["field_ratings"]["red"].append("audit_rights")
                    print(f"[DEBUG] Post-process: audit_rights -> RED (no audit provision)")

        else:
            # Ensure field_ratings exists with proper structure
            parsed["field_ratings"] = {"green": [], "yellow": [], "red": []}
            print("[DEBUG] No field_ratings found, initialized empty")

        # Store full text for re-parsing
        if text:
            parsed["full_text"] = text
            parsed["text_preview"] = text[:500]
        elif use_pdf_fallback:
            # PDF was parsed directly - mark it so reparse knows to require re-upload
            parsed["pdf_parsed_directly"] = True
            parsed["full_text"] = None
            parsed["text_preview"] = "[PDF parsed directly by AI - text extraction unavailable]"
        else:
            parsed["full_text"] = None
            parsed["text_preview"] = None

        # Include extraction metadata in the result for frontend access
        if extraction_metadata:
            parsed["extraction_metadata"] = extraction_metadata

            # Set extraction_issue flag based on quality score
            quality_score = extraction_metadata.get("quality_score", 100)
            warnings = extraction_metadata.get("warnings", [])

            # Flag extraction issues for low-quality extractions
            if quality_score < 50:
                parsed["extraction_issue"] = True
                parsed["extraction_issue_severity"] = "high" if quality_score < 30 else "medium"
                parsed["extraction_issue_message"] = (
                    "Text extraction quality was poor. Some information may be missing or inaccurate. "
                    "Consider re-extracting with Vision API for scanned/image PDFs."
                )
                # Add extraction quality warning to general_assessment if present
                if "general_assessment" in parsed and isinstance(parsed["general_assessment"], str):
                    parsed["general_assessment"] = (
                        f"⚠️ Note: Document extraction quality was {quality_score}/100. "
                        f"Some terms may be missing or incomplete. {parsed['general_assessment']}"
                    )
            elif quality_score < 70 and warnings:
                parsed["extraction_issue"] = True
                parsed["extraction_issue_severity"] = "low"
                parsed["extraction_issue_message"] = (
                    f"Some extraction warnings detected: {'; '.join(warnings[:2])}"
                )

        print(f"[DEBUG] Successfully parsed agreement with {len(parsed)} fields")
        return parsed

    except json.JSONDecodeError as e:
        print(f"[ERROR] Parsing AI response as JSON: {e}")
        print(f"[ERROR] Response was: {response_text[:500] if 'response_text' in dir() else 'N/A'}")
        return {"text_preview": text[:1000] if text else None, "parse_error": "Failed to parse AI response"}
    except Exception as e:
        import traceback
        print(f"[ERROR] Calling AI for agreement parsing: {e}")
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        return {"text_preview": text[:1000] if text else None, "parse_error": str(e)}


@agreements_router.get("", response_model=AgreementListResponse)
async def get_agreements(
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Get all agreements for the current user with analysis summaries"""
    agreements = db.query(Agreement).filter(Agreement.user_id == user.id).all()

    agreement_responses = []
    for a in agreements:
        # Extract summary fields from parsed_content if available
        overall_rating = None
        red_flag_count = None
        critical_flags_count = None
        high_flags_count = None
        medium_flags_count = None

        extraction_quality_score = None
        extraction_method = None

        if a.parsed_content and isinstance(a.parsed_content, dict):
            # Try new analyzer format first
            overall_assessment = a.parsed_content.get("overall_assessment", {})
            if isinstance(overall_assessment, dict):
                overall_rating = overall_assessment.get("rating")
                critical_flags_count = overall_assessment.get("critical_flags", 0)
                high_flags_count = overall_assessment.get("high_flags", 0)
                medium_flags_count = overall_assessment.get("medium_flags", 0)

            # Count red flags from red_flags array
            red_flags_array = a.parsed_content.get("red_flags", [])
            if isinstance(red_flags_array, list):
                red_flag_count = len(red_flags_array)

            # Fallback to legacy format if new format not found
            if red_flag_count is None:
                field_ratings = a.parsed_content.get("field_ratings", {})
                if isinstance(field_ratings, dict):
                    red_flag_count = len(field_ratings.get("red", []))

            # Extract extraction quality metadata
            extraction_metadata = a.parsed_content.get("extraction_metadata", {})
            if isinstance(extraction_metadata, dict):
                extraction_quality_score = extraction_metadata.get("quality_score")
                extraction_method = extraction_metadata.get("method")

        agreement_responses.append(
            AgreementResponse(
                id=a.id,
                filename=a.filename,
                original_filename=a.original_filename,
                file_size=a.file_size,
                agreement_type=str(a.agreement_type.value),
                uploaded_at=a.uploaded_at.isoformat(),
                overall_rating=overall_rating,
                red_flag_count=red_flag_count,
                critical_flags_count=critical_flags_count,
                high_flags_count=high_flags_count,
                medium_flags_count=medium_flags_count,
                extraction_quality_score=extraction_quality_score,
                extraction_method=extraction_method,
            )
        )

    return AgreementListResponse(
        agreements=agreement_responses,
        total=len(agreements),
    )


@agreements_router.get("/{agreement_id}", response_model=AgreementDetailResponse)
async def get_agreement(
    agreement_id: int,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Get a single agreement with all details"""
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    return AgreementDetailResponse(
        id=agreement.id,
        filename=agreement.filename,
        original_filename=agreement.original_filename,
        file_size=agreement.file_size,
        agreement_type=str(agreement.agreement_type.value),
        uploaded_at=agreement.uploaded_at.isoformat(),
        parsed_content=agreement.parsed_content,
        analysis_version=agreement.analysis_version,
    )


VALID_EXTRACTION_METHODS = {"auto", "vision", "standard"}


@agreements_router.post("", response_model=AgreementResponse)
async def upload_agreement(
    file: UploadFile,
    extraction_method: str = Form(default="auto"),
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Upload and parse an agreement document

    Args:
        file: The agreement document (PDF, DOC, DOCX)
        extraction_method: Text extraction method - "auto" (default), "vision", or "standard"
    """
    # Validate extraction_method
    if extraction_method not in VALID_EXTRACTION_METHODS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid extraction_method. Must be one of: {', '.join(VALID_EXTRACTION_METHODS)}",
        )

    # Validate file type
    allowed_extensions = [".pdf", ".doc", ".docx"]
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(allowed_extensions)}",
        )

    temp = NamedTemporaryFile(delete=False, suffix=file_ext)
    temp_path = temp.name

    try:
        # Save file temporarily
        contents = await file.read()
        
        # Debug: Log received file info
        print(f"[DEBUG] Received file: {file.filename}, size: {len(contents)} bytes")
        if contents[:8]:
            header = contents[:8]
            print(f"[DEBUG] File header (hex): {header.hex()}")
            print(f"[DEBUG] File header (ascii): {header[:8]}")
            is_pdf = contents[:4] == b'%PDF'
            print(f"[DEBUG] Is valid PDF header: {is_pdf}")
        
        with temp as f:
            f.write(contents)
        
        # Verify file was written correctly
        import os
        written_size = os.path.getsize(temp_path)
        print(f"[DEBUG] Written to temp file: {temp_path}, size on disk: {written_size} bytes")

        # Extract text based on file type with enhanced extraction
        extraction_result = None
        if file_ext == ".pdf":
            extraction_result = extract_text_from_pdf(temp_path, preferred_method=extraction_method)
            text = extraction_result.text
        elif file_ext in [".doc", ".docx"]:
            extraction_result = extract_text_from_docx(temp_path)
            text = extraction_result.text
        else:
            text = ""

        # Build extraction metadata for AI analyzer
        extraction_metadata = None
        if extraction_result:
            extraction_metadata = {
                "method": extraction_result.method,
                "quality_score": extraction_result.quality_score,
                "warnings": extraction_result.warnings,
                "character_count": extraction_result.character_count,
                "page_count": extraction_result.page_count
            }
            print(f"[DEBUG] Extraction complete: method={extraction_result.method}, quality={extraction_result.quality_score}")

        # Detect agreement type using AI analysis of content
        agreement_type = detect_agreement_type(text, file.filename, file_path=temp_path)

        # Parse agreement with AI to extract structured data
        # Pass file_path so Claude can read PDF directly if text extraction failed
        # Pass extraction_metadata so AI can adjust its analysis accordingly
        parsed_content = parse_agreement_with_ai(
            text,
            str(agreement_type.value),
            file_path=temp_path,
            extraction_metadata=extraction_metadata
        )

        # Generate unique filename for storage
        unique_filename = f"{uuid.uuid4()}{file_ext}"

        # Save file to persistent storage for future re-extraction
        file_storage_path = None
        try:
            file_storage_path = save_agreement_file(contents, unique_filename)
            logger.info(f"Saved agreement file to persistent storage: {file_storage_path}")
        except Exception as storage_error:
            logger.warning(f"Failed to save agreement to persistent storage: {storage_error}")
            # Continue without persistent storage - re-extraction won't be available

        # Create database record with extraction metadata
        agreement = Agreement(
            user_id=user.id,
            filename=unique_filename,
            original_filename=file.filename,
            file_size=len(contents),
            agreement_type=agreement_type,
            parsed_content=parsed_content,
            analysis_version="1.0",  # Track analyzer version for schema evolution
            # Populate extraction metadata columns from ExtractionResult
            extraction_method=extraction_metadata.get("method") if extraction_metadata else None,
            extraction_quality_score=extraction_metadata.get("quality_score", 0) if extraction_metadata else 0,
            extraction_warnings=extraction_metadata.get("warnings", []) if extraction_metadata else [],
            text_character_count=extraction_metadata.get("character_count", 0) if extraction_metadata else 0,
            # Persistent file storage path for re-extraction
            file_storage_path=file_storage_path,
        )

        db.add(agreement)
        db.commit()
        db.refresh(agreement)

        # Extract summary fields from parsed_content for response
        overall_rating = None
        red_flag_count = None
        critical_flags_count = None
        high_flags_count = None
        medium_flags_count = None

        if parsed_content and isinstance(parsed_content, dict):
            # New analyzer format
            overall_assessment = parsed_content.get("overall_assessment", {})
            if isinstance(overall_assessment, dict):
                overall_rating = overall_assessment.get("rating")
                critical_flags_count = overall_assessment.get("critical_flags", 0)
                high_flags_count = overall_assessment.get("high_flags", 0)
                medium_flags_count = overall_assessment.get("medium_flags", 0)

            # Count red flags from red_flags array
            red_flags_array = parsed_content.get("red_flags", [])
            if isinstance(red_flags_array, list):
                red_flag_count = len(red_flags_array)

            # Fallback to legacy format if new format not found
            if red_flag_count is None:
                field_ratings = parsed_content.get("field_ratings", {})
                if isinstance(field_ratings, dict):
                    red_flag_count = len(field_ratings.get("red", []))

        # Create notification for agreement parsed
        try:
            notification_service = NotificationService(db)
            notification_service.create_agreement_parsed_notification(
                user_id=user.id,
                agreement_id=agreement.id,
                filename=file.filename,
                agreement_type=str(agreement_type.value),
                red_flags=red_flag_count or 0,
            )
        except Exception as notif_error:
            print(f"[WARN] Failed to create agreement notification: {notif_error}")

        # Create notifications for critical red flags (RF01-RF06)
        try:
            CRITICAL_FLAG_IDS = ["RF01", "RF02", "RF03", "RF04", "RF05", "RF06"]
            red_flags_array = parsed_content.get("red_flags", []) if parsed_content else []
            if isinstance(red_flags_array, list):
                for flag in red_flags_array:
                    if isinstance(flag, dict) and flag.get("id") in CRITICAL_FLAG_IDS:
                        try:
                            notification_service.create_critical_red_flag_notification(
                                user_id=user.id,
                                agreement_id=agreement.id,
                                filename=file.filename,
                                red_flag=flag,
                            )
                        except Exception as flag_notif_error:
                            print(f"[WARN] Failed to create critical red flag notification for {flag.get('id')}: {flag_notif_error}")
        except Exception as critical_flags_error:
            print(f"[WARN] Failed to process critical red flags: {critical_flags_error}")

        return AgreementResponse(
            id=agreement.id,
            filename=agreement.filename,
            original_filename=agreement.original_filename,
            file_size=agreement.file_size,
            agreement_type=str(agreement.agreement_type.value),
            uploaded_at=agreement.uploaded_at.isoformat(),
            overall_rating=overall_rating,
            red_flag_count=red_flag_count,
            critical_flags_count=critical_flags_count,
            high_flags_count=high_flags_count,
            medium_flags_count=medium_flags_count,
            # Use model columns for extraction metadata
            extraction_quality_score=agreement.extraction_quality_score,
            extraction_method=agreement.extraction_method,
        )

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    finally:
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass


@agreements_router.patch("/{agreement_id}", response_model=AgreementDetailResponse)
async def update_agreement(
    agreement_id: int,
    update_data: AgreementUpdateRequest,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Update agreement details including parsed content"""
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    # Update agreement type if provided
    if update_data.agreement_type:
        type_map = {
            "producer agreement": AgreementType.PRODUCER_AGREEMENT,
            "publishing": AgreementType.PUBLISHING,
            "management": AgreementType.MANAGEMENT,
        }
        new_type = type_map.get(update_data.agreement_type.lower())
        if new_type:
            agreement.agreement_type = new_type

    # Update parsed content if provided
    if update_data.parsed_content is not None:
        # Merge with existing parsed_content
        existing = agreement.parsed_content or {}
        existing.update(update_data.parsed_content)
        agreement.parsed_content = existing

    db.commit()
    db.refresh(agreement)

    return AgreementDetailResponse(
        id=agreement.id,
        filename=agreement.filename,
        original_filename=agreement.original_filename,
        file_size=agreement.file_size,
        agreement_type=str(agreement.agreement_type.value),
        uploaded_at=agreement.uploaded_at.isoformat(),
        parsed_content=agreement.parsed_content,
        analysis_version=agreement.analysis_version,
    )


@agreements_router.delete("/{agreement_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_agreement(
    agreement_id: int,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Delete an agreement"""
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    # Clean up stored file if exists
    if agreement.file_storage_path:
        delete_agreement_file(agreement.file_storage_path)

    db.delete(agreement)
    db.commit()


@agreements_router.post("/{agreement_id}/reparse", response_model=AgreementDetailResponse)
async def reparse_agreement(
    agreement_id: int,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Re-parse an existing agreement with AI to extract structured data"""
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    # Get the full text from existing parsed_content (fall back to text_preview for older agreements)
    existing_content = agreement.parsed_content or {}
    text = existing_content.get("full_text") or existing_content.get("text_preview", "")

    if not text:
        raise HTTPException(
            status_code=400,
            detail="No document text available for re-parsing. Please re-upload the document.",
        )

    # Re-parse with AI
    parsed_content = parse_agreement_with_ai(text, str(agreement.agreement_type.value))

    # Update the agreement with new analysis and version
    agreement.parsed_content = parsed_content
    agreement.analysis_version = "1.0"  # Update version on reparse
    db.commit()
    db.refresh(agreement)

    return AgreementDetailResponse(
        id=agreement.id,
        filename=agreement.filename,
        original_filename=agreement.original_filename,
        file_size=agreement.file_size,
        agreement_type=str(agreement.agreement_type.value),
        uploaded_at=agreement.uploaded_at.isoformat(),
        parsed_content=agreement.parsed_content,
        analysis_version=agreement.analysis_version,
    )


class ReExtractRequest(BaseModel):
    extraction_method: str = "vision"  # "auto", "vision", or "standard"


class ReExtractResponse(BaseModel):
    """Response for re-extraction with quality comparison."""
    id: int
    filename: str
    original_filename: str
    file_size: int
    agreement_type: str
    uploaded_at: str
    parsed_content: Optional[EnhancedAgreementAnalysis] = None
    analysis_version: Optional[str] = None
    # Re-extraction specific fields
    old_quality_score: int
    new_quality_score: int
    quality_improvement: int
    old_extraction_method: Optional[str] = None
    new_extraction_method: str

    class Config:
        from_attributes = True


@agreements_router.post("/{agreement_id}/re-extract", response_model=ReExtractResponse)
async def re_extract_agreement(
    agreement_id: int,
    request: ReExtractRequest,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """
    Re-extract text from an agreement using a different extraction method.

    This endpoint performs TRUE re-extraction from the original file stored during upload,
    not just re-analysis of cached text. This addresses the root cause of poor initial
    text extraction leading to excessive 'Not found in document' results.

    The multi-layer extraction strategy ensures optimal results:
    - "vision": Claude PDF Vision API (best for scanned PDFs, complex layouts, tables)
    - "auto": Intelligent fallback (pdfplumber -> vision if quality < 50 -> pypdf2)
    - "standard": pdfplumber/pypdf2 only (faster, no API calls, for simple PDFs)

    extraction_method options:
    - "auto": Try best method automatically (pdfplumber -> vision -> pypdf2)
    - "vision": Force Claude PDF Vision API (best for scanned/complex PDFs)
    - "standard": Use pdfplumber/pypdf2 only (faster, no API calls)

    Returns quality comparison metrics to show improvement from re-extraction.
    """
    # Validate extraction method
    if request.extraction_method not in VALID_EXTRACTION_METHODS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid extraction_method. Must be one of: {', '.join(VALID_EXTRACTION_METHODS)}",
        )

    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    # Record old extraction quality for comparison
    old_quality_score = agreement.extraction_quality_score or 0
    old_extraction_method = agreement.extraction_method

    # Check if we have the original file stored for re-extraction
    if not agreement.file_storage_path:
        raise HTTPException(
            status_code=400,
            detail="Original file not available for re-extraction. This agreement was uploaded before persistent storage was enabled. Please re-upload the document to enable re-extraction.",
        )

    # Get the file path and validate
    file_path = get_agreement_file_path(agreement.file_storage_path)
    validation = validate_file_for_reextraction(file_path, agreement.file_size)

    if not validation["valid"]:
        raise HTTPException(
            status_code=400,
            detail=validation["error"],
        )

    # Determine file type from stored filename
    file_ext = os.path.splitext(agreement.filename)[1].lower()

    logger.info(
        f"Starting re-extraction for agreement {agreement_id}: "
        f"method={request.extraction_method}, old_quality={old_quality_score}, "
        f"file_type={file_ext}, file_path={file_path}"
    )

    # Perform actual re-extraction from original file
    extraction_result = None
    try:
        if file_ext == ".pdf":
            extraction_result = extract_text_from_pdf(
                str(file_path),
                preferred_method=request.extraction_method
            )
        elif file_ext in [".doc", ".docx"]:
            extraction_result = extract_text_from_docx(str(file_path))
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type for re-extraction: {file_ext}",
            )
    except Exception as e:
        logger.error(f"Re-extraction failed for agreement {agreement_id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Re-extraction failed: {str(e)}",
        )

    if not extraction_result or not extraction_result.text:
        raise HTTPException(
            status_code=500,
            detail="Re-extraction produced no text. The file may be corrupted or unsupported.",
        )

    # Build extraction metadata
    extraction_metadata = {
        "method": extraction_result.method,
        "quality_score": extraction_result.quality_score,
        "warnings": extraction_result.warnings,
        "character_count": extraction_result.character_count,
        "page_count": extraction_result.page_count,
        "re_extracted": True,
        "previous_method": old_extraction_method,
        "previous_quality": old_quality_score,
    }

    new_quality_score = extraction_result.quality_score
    quality_improvement = new_quality_score - old_quality_score

    logger.info(
        f"Re-extraction complete for agreement {agreement_id}: "
        f"new_method={extraction_result.method}, new_quality={new_quality_score}, "
        f"improvement={quality_improvement:+d}"
    )

    # Optionally re-detect agreement type if extraction quality improved significantly
    agreement_type = agreement.agreement_type
    if quality_improvement > 20 or old_quality_score < 50:
        try:
            new_type = detect_agreement_type(
                extraction_result.text,
                agreement.original_filename,
                file_path=str(file_path)
            )
            if new_type != agreement_type:
                logger.info(
                    f"Agreement type changed from {agreement_type.value} to {new_type.value} "
                    f"after re-extraction"
                )
                agreement_type = new_type
        except Exception as type_error:
            logger.warning(f"Failed to re-detect agreement type: {type_error}")

    # Re-parse with AI using fresh extraction and metadata
    parsed_content = parse_agreement_with_ai(
        extraction_result.text,
        str(agreement_type.value),
        file_path=str(file_path),
        extraction_metadata=extraction_metadata
    )

    # Update the agreement with new analysis and extraction metadata
    agreement.parsed_content = parsed_content
    agreement.agreement_type = agreement_type
    agreement.analysis_version = "2.0"  # Bump version for re-extraction
    # Update extraction metadata columns
    agreement.extraction_method = extraction_result.method
    agreement.extraction_quality_score = new_quality_score
    agreement.extraction_warnings = extraction_result.warnings
    agreement.text_character_count = extraction_result.character_count

    db.commit()
    db.refresh(agreement)

    # Log monitoring metrics
    logger.info(
        f"[MONITORING] Re-extraction success: agreement_id={agreement_id}, "
        f"user_id={user.id}, old_score={old_quality_score}, new_score={new_quality_score}, "
        f"improvement={quality_improvement}, method={extraction_result.method}"
    )

    return ReExtractResponse(
        id=agreement.id,
        filename=agreement.filename,
        original_filename=agreement.original_filename,
        file_size=agreement.file_size,
        agreement_type=str(agreement.agreement_type.value),
        uploaded_at=agreement.uploaded_at.isoformat(),
        parsed_content=agreement.parsed_content,
        analysis_version=agreement.analysis_version,
        old_quality_score=old_quality_score,
        new_quality_score=new_quality_score,
        quality_improvement=quality_improvement,
        old_extraction_method=old_extraction_method,
        new_extraction_method=extraction_result.method,
    )


class ChangeTypeRequest(BaseModel):
    new_type: str  # "producer agreement", "publishing", or "management"


@agreements_router.post("/{agreement_id}/change-type", response_model=AgreementDetailResponse)
async def change_agreement_type(
    agreement_id: int,
    request: ChangeTypeRequest,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """Change agreement type and re-parse with the correct prompt for that type"""
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    # Map string to AgreementType enum
    type_map = {
        "producer agreement": AgreementType.PRODUCER_AGREEMENT,
        "publishing": AgreementType.PUBLISHING,
        "management": AgreementType.MANAGEMENT,
    }
    new_type = type_map.get(request.new_type.lower())
    if not new_type:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid agreement type. Must be one of: {', '.join(type_map.keys())}",
        )

    # Get the full text from existing parsed_content
    existing_content = agreement.parsed_content or {}
    text = existing_content.get("full_text") or existing_content.get("text_preview", "")

    if not text:
        raise HTTPException(
            status_code=400,
            detail="No document text available for re-parsing. Please re-upload the document.",
        )

    # Update agreement type
    agreement.agreement_type = new_type

    # Re-parse with AI using the NEW type's prompt
    parsed_content = parse_agreement_with_ai(text, str(new_type.value))

    # Update the agreement with new analysis and version
    agreement.parsed_content = parsed_content
    agreement.analysis_version = "1.0"  # Update version on type change
    db.commit()
    db.refresh(agreement)

    return AgreementDetailResponse(
        id=agreement.id,
        filename=agreement.filename,
        original_filename=agreement.original_filename,
        file_size=agreement.file_size,
        agreement_type=str(agreement.agreement_type.value),
        uploaded_at=agreement.uploaded_at.isoformat(),
        parsed_content=agreement.parsed_content,
        analysis_version=agreement.analysis_version,
    )


@agreements_router.get("/{agreement_id}/report")
async def download_agreement_report(
    agreement_id: int,
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """
    Generate and download a DOCX report for an agreement.

    Returns a professionally formatted Word document with:
    - Executive summary with color-coded rating
    - Red flag analysis (for producer agreements using new analyzer)
    - Terms analysis tables with color-coded assessments
    - Negotiation priorities
    - Financial projections (if available)
    """
    agreement = (
        db.query(Agreement)
        .filter(Agreement.id == agreement_id, Agreement.user_id == user.id)
        .first()
    )

    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    if not agreement.parsed_content:
        raise HTTPException(
            status_code=400,
            detail="Agreement has not been analyzed yet. Please wait for analysis to complete.",
        )

    try:
        # Generate DOCX report
        generator = DocumentGenerator()
        buffer = generator.generate(agreement.parsed_content)

        # Create filename from original filename
        base_name = os.path.splitext(agreement.original_filename)[0]
        download_filename = f"{base_name}_analysis.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{download_filename}"'
            },
        )

    except Exception as e:
        print(f"[ERROR] Failed to generate agreement report: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate report: {str(e)}",
        )


class ExtractionStatsResponse(BaseModel):
    """Response model for extraction quality statistics"""
    total_agreements: int
    average_quality_score: float
    method_distribution: dict  # {method: count}
    quality_distribution: dict  # {quality_level: count}
    low_quality_count: int  # Agreements with score < 50
    agreements_needing_reextraction: List[dict]  # List of {id, filename, score}


@agreements_router.get("/stats/extraction", response_model=ExtractionStatsResponse)
async def get_extraction_stats(
    user: User = Depends(get_user),
    db: Session = Depends(get_session),
):
    """
    Get extraction quality statistics for all user agreements.

    Returns:
    - Total agreements count
    - Average extraction quality score
    - Distribution by extraction method
    - Distribution by quality level (low/moderate/good)
    - List of agreements that may need re-extraction (score < 50)
    """
    from sqlalchemy import func

    agreements = db.query(Agreement).filter(Agreement.user_id == user.id).all()

    if not agreements:
        return ExtractionStatsResponse(
            total_agreements=0,
            average_quality_score=0.0,
            method_distribution={},
            quality_distribution={"low": 0, "moderate": 0, "good": 0},
            low_quality_count=0,
            agreements_needing_reextraction=[],
        )

    # Calculate statistics
    total = len(agreements)
    quality_scores = []
    method_counts = {}
    quality_levels = {"low": 0, "moderate": 0, "good": 0}
    needs_reextraction = []

    for a in agreements:
        # Get quality score from parsed_content.extraction_metadata
        score = None
        method = None

        if a.parsed_content:
            metadata = a.parsed_content.get("extraction_metadata", {})
            if isinstance(metadata, dict):
                score = metadata.get("quality_score")
                method = metadata.get("method")

        if score is not None:
            quality_scores.append(score)

            # Categorize by quality level
            if score < 50:
                quality_levels["low"] += 1
                needs_reextraction.append({
                    "id": a.id,
                    "filename": a.original_filename,
                    "score": score,
                    "method": method,
                })
            elif score < 70:
                quality_levels["moderate"] += 1
            else:
                quality_levels["good"] += 1

        # Count methods
        if method:
            method_counts[method] = method_counts.get(method, 0) + 1

    avg_score = sum(quality_scores) / len(quality_scores) if quality_scores else 0.0

    return ExtractionStatsResponse(
        total_agreements=total,
        average_quality_score=round(avg_score, 1),
        method_distribution=method_counts,
        quality_distribution=quality_levels,
        low_quality_count=quality_levels["low"],
        agreements_needing_reextraction=needs_reextraction[:20],  # Limit to 20
    )
