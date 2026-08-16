"""
Document Generator Service
Generates professional DOCX reports from music agreement analysis.
Uses python-docx with styled tables and color-coded assessments.
"""

import io
import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# Color definitions ported from MusicAgreementAnalyzer.js (lines 16-21)
COLORS = {
    "RED": {"bg": "FFEBEE", "border": "EF9A9A", "text": "C62828", "label": "Unfavorable"},
    "YELLOW": {"bg": "FFF8E1", "border": "FFE082", "text": "F57F17", "label": "Neutral"},
    "GREEN": {"bg": "E8F5E9", "border": "A5D6A7", "text": "2E7D32", "label": "Favorable"},
    "GRAY": {"bg": "F5F5F5", "border": "E0E0E0", "text": "757575", "label": "Not Found"},
}

SECTION_TITLES = {
    "financial": "1.1 Financial Terms",
    "rights": "1.2 Rights Granted",
    "credit": "1.3 Credit & Attribution",
    "legal": "1.4 Legal Protections",
    "administrative": "1.5 Administrative Terms",
    "publishing": "1.6 Composition & Publishing",
}

SECTION_ORDER = ["financial", "rights", "credit", "legal", "administrative", "publishing"]


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor object."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _set_cell_shading(cell, fill_color: str):
    """Set background shading for a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _format_field_name(name: str) -> str:
    """Convert snake_case field name to Title Case."""
    return name.replace("_", " ").title()


def _create_element(name: str) -> OxmlElement:
    """Create an OxmlElement with the given name."""
    return OxmlElement(name)


def _create_attribute(element: OxmlElement, name: str, value: str):
    """Set an XML attribute with qualified name on an element."""
    element.set(qn(name), value)


def _add_page_number(run):
    """Insert PAGE field XML into a run element for current page number."""
    fld_char_begin = _create_element("w:fldChar")
    _create_attribute(fld_char_begin, "w:fldCharType", "begin")

    instr_text = _create_element("w:instrText")
    _create_attribute(instr_text, "xml:space", "preserve")
    instr_text.text = " PAGE "

    fld_char_separate = _create_element("w:fldChar")
    _create_attribute(fld_char_separate, "w:fldCharType", "separate")

    fld_char_end = _create_element("w:fldChar")
    _create_attribute(fld_char_end, "w:fldCharType", "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def _add_total_pages(run):
    """Insert NUMPAGES field XML into a run element for total page count."""
    fld_char_begin = _create_element("w:fldChar")
    _create_attribute(fld_char_begin, "w:fldCharType", "begin")

    instr_text = _create_element("w:instrText")
    _create_attribute(instr_text, "xml:space", "preserve")
    instr_text.text = " NUMPAGES "

    fld_char_separate = _create_element("w:fldChar")
    _create_attribute(fld_char_separate, "w:fldCharType", "separate")

    fld_char_end = _create_element("w:fldChar")
    _create_attribute(fld_char_end, "w:fldCharType", "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


class DocumentGenerator:
    """
    Generates professional DOCX reports from agreement analysis data.
    Includes color-coded tables, red flag sections, and negotiation priorities.
    """

    def __init__(self):
        """Initialize the document generator."""
        pass

    def generate(self, analysis: Dict[str, Any]) -> io.BytesIO:
        """
        Generate a DOCX report from analysis data.

        Args:
            analysis: The analysis result from MusicAgreementAnalyzer.

        Returns:
            BytesIO buffer containing the DOCX file.
        """
        doc = Document()

        # Set document styles
        self._setup_styles(doc)

        # Extract data from analysis
        agreement = analysis.get("agreement", {})
        overall = analysis.get("overall_assessment", {})
        terms = analysis.get("terms", {})
        red_flags = analysis.get("red_flags", [])
        negotiation_priorities = analysis.get("negotiation_priorities", [])

        # Determine rating color
        rating = overall.get("rating", "NEUTRAL")
        rating_color_map = {
            "FAVORABLE": "GREEN",
            "NEUTRAL": "YELLOW",
            "UNFAVORABLE": "RED",
            "HIGHLY_UNFAVORABLE": "RED",
        }
        rating_color = rating_color_map.get(rating, "YELLOW")

        # Title
        title = doc.add_heading("Agreement Analysis", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Agreement type subtitle
        agreement_type = agreement.get("type", "Music Agreement")
        if agreement_type:
            agreement_type = agreement_type.replace("_", " ")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(agreement_type)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Track/project name if available
        parties = agreement.get("parties", {})
        track = parties.get("track_or_project")
        if track:
            track_para = doc.add_paragraph()
            track_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = track_para.add_run(f'Track: "{track}"')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Rating banner
        flag_count = len(red_flags)
        self._add_rating_banner(doc, rating, rating_color, flag_count)

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        summary = overall.get("summary", "")
        if summary:
            doc.add_paragraph(summary)

        # Color counts summary
        red_count = overall.get("red_count", 0)
        yellow_count = overall.get("yellow_count", 0)
        green_count = overall.get("green_count", 0)
        gray_count = overall.get("gray_count", 0)

        counts_para = doc.add_paragraph()
        counts_para.add_run("Term Assessment Summary: ").bold = True
        counts_para.add_run(f"{green_count} Favorable | {yellow_count} Neutral | {red_count} Unfavorable | {gray_count} Not Found")

        # Red Flag Analysis
        if red_flags:
            doc.add_page_break()
            doc.add_heading("Red Flag Analysis", level=1)

            for flag in red_flags:
                self._add_red_flag(doc, flag)

        # Terms Analysis
        doc.add_page_break()
        doc.add_heading("Terms Analysis", level=1)

        for section_key in SECTION_ORDER:
            section_title = SECTION_TITLES.get(section_key, section_key.title())
            section_terms = terms.get(section_key, {})

            doc.add_heading(section_title, level=2)

            if not section_terms or not isinstance(section_terms, dict):
                para = doc.add_paragraph()
                run = para.add_run("No terms found in this section.")
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                continue

            self._add_terms_table(doc, section_terms)

        # Negotiation Priorities
        if negotiation_priorities:
            doc.add_page_break()
            doc.add_heading("Negotiation Priorities", level=1)
            self._add_negotiation_table(doc, negotiation_priorities)

        # Favorable Terms
        favorable_terms = analysis.get("favorable_terms", [])
        if favorable_terms:
            doc.add_heading("Favorable Terms", level=1)
            for fav in favorable_terms:
                para = doc.add_paragraph(style="List Bullet")
                term_name = fav.get("term", "")
                value = fav.get("value", "")
                why = fav.get("why_favorable", "")
                para.add_run(f"{term_name}: ").bold = True
                para.add_run(f"{value} - {why}")

        # Financial Projection
        projection = analysis.get("financial_projection", {})
        if projection and projection.get("scenario"):
            doc.add_heading("Financial Projection", level=1)
            self._add_financial_projection(doc, projection)

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        run = disclaimer.add_run("Disclaimer: ")
        run.font.bold = True
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = disclaimer.add_run(
            "This analysis is for informational purposes only and does not constitute legal advice. "
            "Consult with a qualified entertainment attorney before making decisions based on this analysis."
        )
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Configure page layout, header, and footer
        self._configure_page_margins(doc)
        self._add_header(doc)
        self._add_footer(doc)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(f"Generated DOCX report ({len(red_flags)} red flags, {len(negotiation_priorities)} priorities)")

        return buffer

    def _setup_styles(self, doc: Document):
        """Configure document styles."""
        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)

    def _configure_page_margins(self, doc: Document):
        """Configure page margins to 1 inch on all sides."""
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _add_header(self, doc: Document):
        """Add header with confidentiality notice to document."""
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # Get or create paragraph in header
        if header.paragraphs:
            paragraph = header.paragraphs[0]
        else:
            paragraph = header.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run("Agreement Analysis - Confidential")
        run.font.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_footer(self, doc: Document):
        """Add footer with page numbers (Page X of Y) to document."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Get or create paragraph in footer
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Page " text
        run = paragraph.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Add current page number
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _add_page_number(run)

        # Add " of " text
        run = paragraph.add_run(" of ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Add total pages number
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _add_total_pages(run)

    def _add_rating_banner(self, doc: Document, rating: str, color_key: str, flag_count: int):
        """Add a colored rating banner paragraph."""
        color_def = COLORS.get(color_key, COLORS["YELLOW"])

        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add shading to paragraph
        pPr = banner._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_def["bg"])
        pPr.append(shd)

        # Banner text
        rating_text = rating.replace("_", " ")
        if flag_count > 0:
            text = f"! {flag_count} RED FLAGS - {rating_text} !"
        else:
            text = rating_text

        run = banner.add_run(text)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _hex_to_rgb(color_def["text"])

    def _add_red_flag(self, doc: Document, flag: Dict[str, Any]):
        """Add a red flag section to the document."""
        severity = flag.get("severity", "MEDIUM")
        flag_id = flag.get("id", "")
        name = flag.get("name", "")
        clause = flag.get("clause", "")
        quote = flag.get("quote", "")
        impact = flag.get("impact", "")
        recommendation = flag.get("recommendation", "")

        # Determine color based on severity
        sev_color = "RED" if severity in ["CRITICAL", "HIGH"] else "YELLOW"
        color_def = COLORS[sev_color]

        # Flag heading
        heading = doc.add_heading(level=3)
        run = heading.add_run(f"{flag_id}: {name}")
        run.font.color.rgb = _hex_to_rgb(color_def["text"])

        # Severity badge
        sev_para = doc.add_paragraph()
        run = sev_para.add_run(f"Severity: {severity}")
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb(color_def["text"])

        # Clause reference
        if clause:
            clause_para = doc.add_paragraph()
            run = clause_para.add_run(f"Clause: {clause}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Quote (if available)
        if quote:
            quote_para = doc.add_paragraph()
            # Add shading
            pPr = quote_para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F5F5F5")
            pPr.append(shd)

            run = quote_para.add_run(f'"{quote}"')
            run.font.italic = True
            run.font.size = Pt(9)

        # Impact
        if impact:
            impact_para = doc.add_paragraph()
            impact_para.add_run("Impact: ").bold = True
            impact_para.add_run(impact)

        # Recommendation
        if recommendation:
            rec_para = doc.add_paragraph()
            run = rec_para.add_run("Recommendation: ")
            run.bold = True
            run.font.color.rgb = _hex_to_rgb(COLORS["GREEN"]["text"])
            rec_para.add_run(recommendation)

        # Add spacing
        doc.add_paragraph()

    def _add_terms_table(self, doc: Document, terms: Dict[str, Any], skip_na: bool = True):
        """Add a terms analysis table to the document.

        Args:
            doc: The document to add the table to
            terms: Dictionary of terms to display
            skip_na: If True (default), N/A terms are excluded from the report
        """
        # Filter out N/A terms if skip_na is True
        filtered_terms = {}
        na_count = 0
        for field_name, term in terms.items():
            if not isinstance(term, dict):
                continue
            value = term.get("value", "")
            if skip_na and (value == "N/A" or "not applicable" in str(term.get("assessment", "")).lower()):
                na_count += 1
                continue
            filtered_terms[field_name] = term

        # If all terms are N/A, show a message instead of empty table
        if not filtered_terms:
            para = doc.add_paragraph()
            run = para.add_run(f"All {na_count} terms in this section are not applicable to this agreement type.")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        # Create table with header row
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        headers = ["Term", "Value", "Standard", "Assessment"]
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            # Style header cell
            _set_cell_shading(cell, "1A237E")  # Navy blue
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)

        # Data rows
        for field_name, term in filtered_terms.items():
            row = table.add_row()
            cells = row.cells

            # Term name
            cells[0].text = _format_field_name(field_name)
            for run in cells[0].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)

            # Value
            value = term.get("value", "")
            if value == "NOT_FOUND":
                value = "Not specified"
            elif value == "N/A":
                value = "N/A"
            cells[1].text = str(value)[:100]  # Truncate long values
            for run in cells[1].paragraphs[0].runs:
                run.font.size = Pt(9)

            # Industry standard
            cells[2].text = str(term.get("industry_standard", ""))[:80]
            for run in cells[2].paragraphs[0].runs:
                run.font.size = Pt(9)

            # Assessment (colored)
            color = term.get("color", "GRAY")
            color_def = COLORS.get(color, COLORS["GRAY"])
            assessment = term.get("assessment", "")

            cells[3].text = str(assessment)[:120]
            _set_cell_shading(cells[3], color_def["bg"])
            for run in cells[3].paragraphs[0].runs:
                run.font.color.rgb = _hex_to_rgb(color_def["text"])
                run.font.size = Pt(9)

        # Add note about hidden N/A terms
        if skip_na and na_count > 0:
            para = doc.add_paragraph()
            run = para.add_run(f"Note: {na_count} term(s) not applicable to this agreement type were excluded.")
            run.font.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(1.8)
            row.cells[2].width = Inches(1.8)
            row.cells[3].width = Inches(2.2)

    def _add_negotiation_table(self, doc: Document, priorities: List[Dict[str, Any]]):
        """Add a negotiation priorities table to the document."""
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        headers = ["#", "Issue", "Current", "Target", "Impact"]
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            _set_cell_shading(cell, "1A237E")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)

        # Data rows
        for item in priorities:
            row = table.add_row()
            cells = row.cells

            priority = item.get("priority", "")
            impact = item.get("impact", "MEDIUM")

            # Priority number (colored)
            cells[0].text = str(priority)
            color_key = "RED" if impact == "CRITICAL" else "YELLOW"
            _set_cell_shading(cells[0], COLORS[color_key]["bg"])
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[0].paragraphs[0].runs:
                run.font.color.rgb = _hex_to_rgb(COLORS[color_key]["text"])
                run.font.bold = True
                run.font.size = Pt(9)

            # Issue
            issue = item.get("issue") or item.get("term", "")
            cells[1].text = str(issue)
            for run in cells[1].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)

            # Current (red)
            cells[2].text = str(item.get("current", ""))[:60]
            _set_cell_shading(cells[2], COLORS["RED"]["bg"])
            for run in cells[2].paragraphs[0].runs:
                run.font.color.rgb = _hex_to_rgb(COLORS["RED"]["text"])
                run.font.size = Pt(9)

            # Target (green)
            cells[3].text = str(item.get("target", ""))[:60]
            _set_cell_shading(cells[3], COLORS["GREEN"]["bg"])
            for run in cells[3].paragraphs[0].runs:
                run.font.color.rgb = _hex_to_rgb(COLORS["GREEN"]["text"])
                run.font.size = Pt(9)

            # Impact
            cells[4].text = str(impact)
            for run in cells[4].paragraphs[0].runs:
                run.font.size = Pt(9)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.4)
            row.cells[1].width = Inches(1.8)
            row.cells[2].width = Inches(1.6)
            row.cells[3].width = Inches(1.6)
            row.cells[4].width = Inches(0.8)

    def _add_financial_projection(self, doc: Document, projection: Dict[str, Any]):
        """Add financial projection section to the document."""
        scenario = projection.get("scenario", "")
        if scenario:
            para = doc.add_paragraph()
            para.add_run("Scenario: ").bold = True
            para.add_run(scenario)

        # Create simple table for projections
        fields = [
            ("Estimated Advance", "estimated_advance"),
            ("Recording Royalties", "estimated_recording_royalties"),
            ("Sync Income", "estimated_sync_income"),
            ("SoundExchange", "estimated_soundexchange"),
            ("Publishing", "estimated_publishing"),
        ]

        for label, key in fields:
            value = projection.get(key)
            if value:
                para = doc.add_paragraph()
                para.add_run(f"{label}: ").bold = True
                para.add_run(str(value))

        # Key insight
        insight = projection.get("key_insight", "")
        if insight:
            doc.add_paragraph()
            para = doc.add_paragraph()
            para.add_run("Key Insight: ").bold = True
            para.add_run(insight)


def generate_agreement_report(analysis: Dict[str, Any]) -> io.BytesIO:
    """
    Convenience function to generate an agreement report.

    Args:
        analysis: The analysis result from MusicAgreementAnalyzer.

    Returns:
        BytesIO buffer containing the DOCX file.
    """
    generator = DocumentGenerator()
    return generator.generate(analysis)
